"""
ESGIntel — Policy Document Generator
=====================================
Produces fully-developed, McKinsey-style company policy documents (.docx) for
the demo company {D.COMPANY['short']} S.p.A.

Public contract:
    build_policy(key: str, out_dir: str) -> str
    key in {'human_rights', 'whistleblower'}
    returns absolute path to the generated .docx file.
"""

import os
import sys
import tempfile
from datetime import date
from pathlib import Path

# ── non-interactive matplotlib ──────────────────────────────────────────────
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyArrowPatch
import numpy as np

# ── python-docx ─────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── company data (single source of truth) ───────────────────────────────────
sys.path.insert(0, os.path.dirname(__file__))
import company_data as cd

# ────────────────────────────────────────────────────────────────────────────
# Shortcuts & colour helpers
# ────────────────────────────────────────────────────────────────────────────
B = cd.BRAND


def _rgb(key: str) -> RGBColor:
    h = B[key]
    r, g, b = cd.hex_rgb(h)
    return RGBColor(r, g, b)


def _hex(key: str) -> str:
    return B[key]


# ════════════════════════════════════════════════════════════════════════════
# DOCX CORE HELPERS
# ════════════════════════════════════════════════════════════════════════════

def _new_doc() -> Document:
    """Create a new document with A4 page size and branded margins."""
    doc = Document()
    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.2)
        section.bottom_margin = Cm(2.2)
    return doc


def _set_para_fmt(para, space_before=0, space_after=4, line_spacing=None,
                  keep_together=False):
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if line_spacing:
        pf.line_spacing = Pt(line_spacing)
    if keep_together:
        pf.keep_together = True


def _add_heading(doc: Document, text: str, level: int = 1,
                 color_key: str = "ink", size: int = None,
                 space_before: int = 14, space_after: int = 6,
                 page_break_before: bool = False) -> None:
    para = doc.add_paragraph()
    if page_break_before:
        para.paragraph_format.page_break_before = True
    run = para.add_run(text)
    sizes = {1: 18, 2: 13, 3: 11, 4: 10}
    run.font.size  = Pt(size or sizes.get(level, 10))
    run.font.bold  = True
    run.font.color.rgb = _rgb(color_key)
    run.font.name  = B["font_head"]
    _set_para_fmt(para, space_before=space_before, space_after=space_after)
    if level <= 2:
        pPr  = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"),  "4" if level == 1 else "2")
        bot.set(qn("w:space"), "4")
        bot.set(qn("w:color"), B["accent"])
        pBdr.append(bot)
        pPr.append(pBdr)


def _add_body(doc: Document, text: str, bold: bool = False,
              italic: bool = False, color_key: str = "slate",
              size: int = 10, space_after: int = 4,
              indent_cm: float = 0) -> None:
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    run.font.color.rgb = _rgb(color_key)
    run.font.name  = B["font_body"]
    _set_para_fmt(para, space_after=space_after)
    if indent_cm:
        para.paragraph_format.left_indent = Cm(indent_cm)


def _add_bullet(doc: Document, text: str, level: int = 0,
                size: int = 10, color_key: str = "slate") -> None:
    para = doc.add_paragraph(style="List Bullet")
    run  = para.add_run(text)
    run.font.size  = Pt(size)
    run.font.color.rgb = _rgb(color_key)
    run.font.name  = B["font_body"]
    _set_para_fmt(para, space_after=2)
    if level:
        para.paragraph_format.left_indent = Cm(level * 0.8)


def _add_caption(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.font.size  = Pt(8)
    run.font.italic = True
    run.font.color.rgb = _rgb("muted")
    _set_para_fmt(para, space_before=2, space_after=8)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ── table helpers ────────────────────────────────────────────────────────────

def _cell_shade(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _cell_borders(cell, color: str = None):
    color = color or B["rule"]
    tc    = cell._tc
    tcPr  = tc.get_or_add_tcPr()
    tcBdr = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBdr.append(el)
    tcPr.append(tcBdr)


def _cell_text(cell, text: str, bold: bool = False, size: int = 9,
               color: str = None, center: bool = False,
               italic: bool = False, wrap: bool = True) -> None:
    color = color or B["slate"]
    para  = cell.paragraphs[0]
    para.clear()
    run   = para.add_run(str(text))
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    run.font.name  = B["font_body"]
    run.font.color.rgb = RGBColor(*cd.hex_rgb(color))
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    para.paragraph_format.left_indent  = Pt(4)


def _make_table(doc: Document, headers: list, col_widths_cm: list,
                header_fill: str = None) -> "Table":
    header_fill = header_fill or B["ink"]
    ncols = len(headers)
    table = doc.add_table(rows=1, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_row = table.rows[0]
    for i, (hdr, w) in enumerate(zip(headers, col_widths_cm)):
        cell = hdr_row.cells[i]
        cell.width = Cm(w)
        _cell_shade(cell, header_fill)
        _cell_borders(cell, B["rule"])
        _cell_text(cell, hdr, bold=True, color="FFFFFF", center=True, size=9)
    return table


def _add_table_row(table, values: list, col_widths_cm: list,
                   fill: str = None, row_idx: int = 0,
                   bold_first: bool = False) -> None:
    fill = fill or (B["band_bg"] if row_idx % 2 == 0 else "FFFFFF")
    row  = table.add_row()
    for i, (val, w) in enumerate(zip(values, col_widths_cm)):
        cell = row.cells[i]
        cell.width = Cm(w)
        _cell_shade(cell, fill)
        _cell_borders(cell, B["rule"])
        _cell_text(cell, str(val) if val is not None else "—",
                   bold=(bold_first and i == 0))


# ── callout / notice box ─────────────────────────────────────────────────────

def _add_callout(doc: Document, text: str, label: str = "NOTICE",
                 fill: str = None, border_color: str = None,
                 label_color: str = None) -> None:
    """Renders a highlighted notice box using a 1-column table."""
    fill         = fill         or "FEF9E7"
    border_color = border_color or B["amber"]
    label_color  = label_color  or B["amber"]
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell  = table.rows[0].cells[0]
    cell.width = Cm(16)
    _cell_shade(cell, fill)
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBdr = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single" if side != "left" else "thick")
        el.set(qn("w:sz"),    "4" if side != "left" else "18")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), border_color)
        tcBdr.append(el)
    tcPr.append(tcBdr)
    # Label line
    p1   = cell.paragraphs[0]
    p1.clear()
    lbl  = p1.add_run(label + "  ")
    lbl.font.bold  = True
    lbl.font.size  = Pt(9)
    lbl.font.name  = B["font_body"]
    lbl.font.color.rgb = RGBColor(*cd.hex_rgb(label_color))
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after  = Pt(0)
    p1.paragraph_format.left_indent  = Pt(6)
    # Body text
    p2   = cell.add_paragraph()
    body = p2.add_run(text)
    body.font.size  = Pt(9)
    body.font.name  = B["font_body"]
    body.font.color.rgb = _rgb("slate")
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(4)
    p2.paragraph_format.left_indent  = Pt(6)
    _set_para_fmt(doc.add_paragraph(), space_after=6)  # spacer after callout


def _add_footer(doc: Document, policy_title: str) -> None:
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para   = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        run    = para.add_run(
            f"ESGIntel — Generated policy template  ·  Legal review required before adoption"
            f"   |   {cd.COMPANY['name']}   |   {policy_title}"
        )
        run.font.size  = Pt(7.5)
        run.font.color.rgb = _rgb("muted")
        run.font.name  = B["font_body"]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr  = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        top  = OxmlElement("w:top")
        top.set(qn("w:val"),   "single")
        top.set(qn("w:sz"),    "2")
        top.set(qn("w:space"), "4")
        top.set(qn("w:color"), B["rule"])
        pBdr.append(top)
        pPr.append(pBdr)


def _add_page_numbers(doc: Document) -> None:
    """Add page numbers to the footer (appended after text)."""
    for section in doc.sections:
        footer = section.footer
        if not footer.paragraphs:
            footer.add_paragraph()
        # Find the existing footer para and append page number field
        para = footer.paragraphs[0]
        # add a tab then PAGE field
        run_tab = para.add_run("   |   Page ")
        run_tab.font.size  = Pt(7.5)
        run_tab.font.color.rgb = _rgb("muted")
        run_tab.font.name  = B["font_body"]
        # PAGE field
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run_pg = para.add_run()
        run_pg._r.append(fldChar1)
        run_pg._r.append(instrText)
        run_pg._r.append(fldChar2)
        run_pg.font.size  = Pt(7.5)
        run_pg.font.color.rgb = _rgb("muted")
        run_pg.font.name  = B["font_body"]


def _insert_image(doc: Document, path: str, width_cm: float = 14,
                  caption: str = "") -> None:
    if not path or not os.path.exists(path):
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run  = para.add_run()
    run.add_picture(path, width=Cm(width_cm))
    if caption:
        _add_caption(doc, caption)


# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════

def _title_page(doc: Document, policy_title: str, subtitle: str,
                version: str = "1.0", status: str = "DRAFT — FOR REVIEW") -> None:
    """Renders a branded title page."""
    # vertical spacer
    for _ in range(6):
        sp = doc.add_paragraph()
        _set_para_fmt(sp, space_after=0)

    # Platform label
    p_plat = doc.add_paragraph()
    r_plat = p_plat.add_run(B["platform"].upper() + "  ·  " + B["tagline"])
    r_plat.font.size  = Pt(9)
    r_plat.font.color.rgb = _rgb("muted")
    r_plat.font.name  = B["font_body"]
    r_plat.font.bold  = True
    p_plat.alignment  = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_fmt(p_plat, space_after=2)

    # Thin accent rule
    p_rule = doc.add_paragraph()
    pPr_r  = p_rule._p.get_or_add_pPr()
    pBdr_r = OxmlElement("w:pBdr")
    bot_r  = OxmlElement("w:bottom")
    bot_r.set(qn("w:val"),   "single")
    bot_r.set(qn("w:sz"),    "6")
    bot_r.set(qn("w:space"), "4")
    bot_r.set(qn("w:color"), B["accent"])
    pBdr_r.append(bot_r)
    pPr_r.append(pBdr_r)
    _set_para_fmt(p_rule, space_after=20)

    # Company name
    p_co = doc.add_paragraph()
    r_co = p_co.add_run(cd.COMPANY["name"])
    r_co.font.size  = Pt(14)
    r_co.font.color.rgb = _rgb("ink")
    r_co.font.name  = B["font_head"]
    r_co.font.bold  = True
    p_co.alignment  = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_fmt(p_co, space_after=8)

    # Policy title
    p_title = doc.add_paragraph()
    r_title = p_title.add_run(policy_title)
    r_title.font.size  = Pt(26)
    r_title.font.color.rgb = _rgb("ink")
    r_title.font.name  = B["font_head"]
    r_title.font.bold  = True
    p_title.alignment  = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_fmt(p_title, space_before=4, space_after=8)

    # Subtitle
    if subtitle:
        p_sub = doc.add_paragraph()
        r_sub = p_sub.add_run(subtitle)
        r_sub.font.size  = Pt(12)
        r_sub.font.color.rgb = _rgb("accent")
        r_sub.font.name  = B["font_head"]
        r_sub.font.italic = True
        p_sub.alignment  = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_fmt(p_sub, space_after=30)

    # Status badge (callout-style table)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_s = tbl.rows[0].cells[0]
    _cell_shade(cell_s, "FEF9E7")
    tc_s   = cell_s._tc
    tcPr_s = tc_s.get_or_add_tcPr()
    tcBdr_s = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), B["amber"])
        tcBdr_s.append(el)
    tcPr_s.append(tcBdr_s)
    ps = cell_s.paragraphs[0]
    ps.clear()
    rs = ps.add_run(status)
    rs.font.size   = Pt(10)
    rs.font.bold   = True
    rs.font.color.rgb = RGBColor(*cd.hex_rgb(B["amber"]))
    rs.font.name   = B["font_body"]
    ps.alignment   = WD_ALIGN_PARAGRAPH.CENTER
    ps.paragraph_format.space_before = Pt(4)
    ps.paragraph_format.space_after  = Pt(4)

    for _ in range(3):
        sp = doc.add_paragraph()
        _set_para_fmt(sp, space_after=0)

    # Version / confidential
    p_conf = doc.add_paragraph()
    r_conf = p_conf.add_run(B["confidential"] + f"   |   Version {version}")
    r_conf.font.size  = Pt(8)
    r_conf.font.color.rgb = _rgb("muted")
    r_conf.font.name  = B["font_body"]
    p_conf.alignment  = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_fmt(p_conf, space_after=2)


def _doc_control_table(doc: Document, version: str, owner: str,
                       approved_by: str, effective_date: str,
                       review_date: str, classification: str) -> None:
    """Document Control table: version, owner, dates, classification."""
    _add_heading(doc, "Document Control", level=2,
                 space_before=8, space_after=6)
    headers = ["Field", "Detail"]
    widths  = [5.5, 10.5]
    table   = _make_table(doc, headers, widths, header_fill=B["slate"])
    rows = [
        ("Document Title",   doc.paragraphs[0].text if False else "—"),
        ("Version",          version),
        ("Status",           "Draft — For Board / Legal Review"),
        ("Prepared by",      f"ESGIntel AI-Assisted ESG Platform  |  {cd.COMPANY['name']} Sustainability Team"),
        ("Policy Owner",     owner),
        ("Approved by",      approved_by),
        ("Effective Date",   effective_date),
        ("Next Review Date", review_date),
        ("Classification",   classification),
        ("Jurisdiction",     f"{cd.COMPANY['domicile']} — Applies to all group entities"),
    ]
    for i, (field, detail) in enumerate(rows):
        _add_table_row(table, [field, detail], widths, row_idx=i, bold_first=True)
    doc.add_paragraph()


# ════════════════════════════════════════════════════════════════════════════
# CHART / VISUAL GENERATORS
# ════════════════════════════════════════════════════════════════════════════

def _chart_hrdd_process(tmpdir: str) -> str:
    """Process-flow bar diagram for HRDD 5-step cycle."""
    steps = [
        "1. Identify\n& Map",
        "2. Assess\nImpacts",
        "3. Integrate\n& Act",
        "4. Track\nProgress",
        "5. Communicate\n& Report",
    ]
    colours = [
        cd.hex_rgb01(B["ink"]),
        cd.hex_rgb01(B["slate"]),
        cd.hex_rgb01(B["accent"]),
        cd.hex_rgb01("1E8E5A"),
        cd.hex_rgb01("C77D11"),
    ]
    fig, ax = plt.subplots(figsize=(7.5, 2.2))
    bar_w   = 0.9
    spacing = 1.6
    for i, (step, col) in enumerate(zip(steps, colours)):
        x = i * spacing
        rect = plt.Rectangle((x, 0.1), bar_w, 0.8,
                              color=col, zorder=2, linewidth=0)
        ax.add_patch(rect)
        ax.text(x + bar_w / 2, 0.5, step,
                ha="center", va="center",
                fontsize=7.5, color="white", fontweight="bold",
                wrap=True, linespacing=1.3)
        if i < len(steps) - 1:
            ax.annotate("", xy=(x + bar_w + 0.28, 0.5),
                        xytext=(x + bar_w + 0.02, 0.5),
                        arrowprops=dict(arrowstyle="->", color="#" + B["muted"],
                                        lw=1.5))
    ax.set_xlim(-0.2, len(steps) * spacing - 0.4)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.set_title("Human Rights Due Diligence Process (UNGP Pillar II)",
                 fontsize=9, color="#" + B["ink"], pad=6, fontweight="bold")
    fig.tight_layout(pad=0.5)
    path = os.path.join(tmpdir, "hrdd_flow.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def _chart_investigation_flow(tmpdir: str) -> str:
    """Process-flow bar diagram for the whistleblower investigation pipeline."""
    steps = [
        "Report\nReceived",
        "Acknowledge\n≤7 days",
        "Initial\nTriage",
        "Investigation\nLaunched",
        "Findings &\nRemediation",
        "Close & Feed-\nback ≤3 months",
    ]
    colours = [
        cd.hex_rgb01(B["ink"]),
        cd.hex_rgb01(B["slate"]),
        cd.hex_rgb01("C77D11"),
        cd.hex_rgb01(B["accent"]),
        cd.hex_rgb01("1E8E5A"),
        cd.hex_rgb01("1F7A5C"),
    ]
    fig, ax = plt.subplots(figsize=(8.5, 2.2))
    bar_w   = 0.9
    spacing = 1.5
    for i, (step, col) in enumerate(zip(steps, colours)):
        x = i * spacing
        rect = plt.Rectangle((x, 0.1), bar_w, 0.8,
                              color=col, zorder=2, linewidth=0)
        ax.add_patch(rect)
        ax.text(x + bar_w / 2, 0.5, step,
                ha="center", va="center",
                fontsize=7, color="white", fontweight="bold",
                linespacing=1.3)
        if i < len(steps) - 1:
            ax.annotate("", xy=(x + bar_w + 0.26, 0.5),
                        xytext=(x + bar_w + 0.02, 0.5),
                        arrowprops=dict(arrowstyle="->", color="#" + B["muted"],
                                        lw=1.5))
    ax.set_xlim(-0.2, len(steps) * spacing - 0.35)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.set_title("Whistleblower Investigation Process & Timeline",
                 fontsize=9, color="#" + B["ink"], pad=6, fontweight="bold")
    fig.tight_layout(pad=0.5)
    path = os.path.join(tmpdir, "wb_flow.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


# ════════════════════════════════════════════════════════════════════════════
# HUMAN RIGHTS POLICY BUILDER
# ════════════════════════════════════════════════════════════════════════════

def _build_human_rights(out_dir: str, tmpdir: str) -> str:
    C   = cd.COMPANY
    doc = _new_doc()

    # ── Title page ──────────────────────────────────────────────────────────
    _title_page(
        doc,
        policy_title="Human Rights Policy",
        subtitle="UNGP-Aligned  ·  CSDDD-Ready  ·  ESRS S2/S3 Referenced",
        version="1.0",
        status="DRAFT — LEGAL & BOARD REVIEW REQUIRED",
    )
    doc.add_page_break()

    # ── Document control ────────────────────────────────────────────────────
    _add_heading(doc, "Document Control & Version History", level=1,
                 page_break_before=False)
    _doc_control_table(
        doc,
        version="1.0  (Initial Draft)",
        owner="Chief Sustainability Officer (CSO)",
        approved_by="Board of Directors — Sustainability & ESG Committee",
        effective_date="[To be completed upon Board approval]",
        review_date="Annually; or upon material change to operations or regulatory landscape",
        classification="Public (post-Board approval)  |  Internal Draft",
    )

    # Compliance urgency callout
    risk_soc01 = next(r for r in cd.RISKS if r["id"] == "SOC-01")
    _add_callout(
        doc,
        text=(
            f"Risk ID SOC-01 — Score {risk_soc01['score']}/100 (High).  "
            f"{risk_soc01['summary']}  |  "
            f"Financial exposure: {risk_soc01['financial']}.  "
            f"Target completion: 12-month implementation plan."
        ),
        label="COMPLIANCE PRIORITY  —  CSDDD 2027 / ESRS S2",
        fill="FAD7D4",
        border_color=B["red"],
        label_color=B["red"],
    )

    doc.add_page_break()

    # ── Section 1: Purpose & Scope ──────────────────────────────────────────
    _add_heading(doc, "1.  Purpose & Scope", level=1)
    _add_body(doc, (
        f"{C['name']} ('{D.COMPANY['short']}' or 'the Company') is a primary steel "
        f"manufacturer (NACE C24) operating across Italy and Spain with approximately "
        f"{C['employees']:,} employees and €{C['revenue_eur_bn']:.2f}bn in annual revenue. "
        f"Our operations — from iron-ore procurement in Brazil and mineral sourcing from "
        f"the Democratic Republic of Congo (DRC) to steelmaking at Taranto, Brescia, Genoa, "
        f"and Bilbao — span complex global value chains with material human rights exposure."
    ), space_after=6)
    _add_body(doc, (
        f"This Policy establishes {D.COMPANY['short']}'s commitment to respecting human rights across "
        "all business activities and relationships. It reflects the Company's responsibility "
        "under internationally recognised frameworks and forms the foundation for our ongoing "
        "Human Rights Due Diligence (HRDD) programme."
    ), space_after=6)

    _add_heading(doc, "1.1  Scope of Application", level=2)
    scope_rows = [
        ("Own Workforce",       f"All {C['employees']:,} employees across Taranto, Brescia, Genoa (Italy) and Bilbao (Spain); includes permanent, fixed-term, agency, and apprenticeship workers."),
        ("Contractors",         "All contractors and sub-contractors performing work on Company premises or on Company-commissioned projects."),
        ("Supply Chain — Tier 1", "All direct suppliers, including iron-ore suppliers (Brazil), mineral sourcing (DRC), energy and logistics providers."),
        ("Supply Chain — Tier 2+", "High-risk sub-tier suppliers as identified through HRDD risk mapping; subject to enhanced due diligence."),
        ("Business Partners",   "Joint venture partners, agents, distributors, and investees over which the Company exercises material influence."),
        ("Communities",         "Local communities adjacent to our operating sites and those affected by our supply-chain activities."),
    ]
    _add_heading(doc, "Scope Matrix", level=3, space_before=6)
    headers = ["Stakeholder Group", "Scope Detail"]
    widths  = [5.0, 11.0]
    tbl     = _make_table(doc, headers, widths)
    for i, (grp, detail) in enumerate(scope_rows):
        _add_table_row(tbl, [grp, detail], widths, row_idx=i, bold_first=True)
    doc.add_paragraph()

    # ── Section 2: Policy Commitment Statement ───────────────────────────────
    _add_heading(doc, "2.  Policy Commitment Statement", level=1,
                 page_break_before=True)
    _add_body(doc, (
        f"{D.COMPANY['short']} is committed to:"
    ), space_after=4)
    commitments = [
        "Respecting all internationally recognised human rights as articulated in the "
        "International Bill of Human Rights and the ILO's Declaration on Fundamental Principles "
        "and Rights at Work.",
        "Conducting ongoing, meaningful Human Rights Due Diligence (HRDD) proportionate to the "
        "scale and context of our operations and value chain.",
        "Providing access to effective remedy for any individuals or communities whose human "
        "rights are adversely impacted by our activities.",
        "Being transparent about our human rights performance, risks, and remediation actions "
        "in line with ESRS S1, S2, and S3 disclosure requirements.",
        "Engaging in genuine dialogue with affected stakeholders — including workers, communities, "
        "civil society organisations, and trade unions (union density: 87% of our workforce).",
        "Integrating human rights considerations into procurement, investment, and partnership decisions.",
        "Meeting or exceeding requirements under the EU Corporate Sustainability Due Diligence "
        "Directive (CSDDD), with full compliance targeted ahead of the 2027 in-scope date.",
    ]
    for c in commitments:
        _add_bullet(doc, c)
    doc.add_paragraph()

    # ── Section 3: Governing Frameworks ────────────────────────────────────
    _add_heading(doc, "3.  Governing Frameworks & Regulatory Context", level=1)
    fw_rows = [
        ("UN Guiding Principles on Business and Human Rights (UNGPs)",
         "2011", "Primary framework. Three pillars: State duty to protect; corporate responsibility to respect; access to remedy."),
        ("ILO Core Conventions (8 fundamental)",
         "Continuous", "Abolition of forced/child labour; non-discrimination; freedom of association; right to bargain collectively."),
        ("OECD Guidelines for Multinational Enterprises",
         "2023 update", "Chapter IV Human Rights: conduct HRDD commensurate with risk."),
        ("EU Corporate Sustainability Due Diligence Directive (CSDDD)",
         "Transposition 2026; in-scope 2027+", "Mandatory HRDD across own operations and supply chain; civil and supervisory liability."),
        ("ESRS S1 — Own Workforce",
         "FY2025 reporting", "Worker conditions, collective bargaining, H&S, pay gaps."),
        ("ESRS S2 — Value-Chain Workers",
         "FY2025 reporting", "Supplier working conditions, HRDD policies, grievance access."),
        ("ESRS S3 — Affected Communities",
         "FY2025 reporting", "Community engagement, impact assessment, land/livelihood rights."),
        ("Italy: D.Lgs. 231/2001",
         "Active", "Organisational model for offence prevention, including labour offences."),
        ("Italy: D.Lgs. 254/2016 (NFI Directive)",
         "Active", "Non-financial information disclosure for large PIEs."),
    ]
    _add_heading(doc, "Regulatory & Framework Reference Table", level=3, space_before=6)
    headers = ["Framework", "Timeline / Status", f"Key Requirement for {D.COMPANY['short']}"]
    widths  = [5.0, 3.5, 7.5]
    tbl     = _make_table(doc, headers, widths)
    for i, row in enumerate(fw_rows):
        _add_table_row(tbl, list(row), widths, row_idx=i)
    doc.add_paragraph()

    # ── Section 4: Salient Human Rights Issues ───────────────────────────────
    _add_heading(doc, "4.  Salient Human Rights Issues", level=1,
                 page_break_before=True)
    _add_body(doc, (
        "Based on our preliminary Human Rights Impact Assessment and value-chain risk mapping, "
        f"the following salient human rights issues have been identified for {D.COMPANY['short']}:"
    ), space_after=6)

    salient_rows = [
        ("HIGH", "Forced & Bonded Labour",
         "DRC mineral supply chain (artisanal mining); Brazilian iron-ore contract workers.",
         "Enhanced supplier audits; contractual prohibitions; SAQ (Tier 1 and high-risk Tier 2)."),
        ("HIGH", "Child Labour",
         "Artisanal and small-scale mining (ASM) in DRC; potential in agricultural land-use in Brazil port logistics.",
         "Zero-tolerance contractual clause; OECD Due Diligence Guidance for Minerals alignment."),
        ("HIGH", "Freedom of Association & Collective Bargaining",
         "Risk in DRC and Brazilian supplier operations where trade union rights are restricted.",
         "Supplier Code of Conduct requirement; supplier dialogue and capacity building."),
        ("HIGH", "Occupational Health & Safety",
         f"Own workforce: LTIFR 2.8 (33% above sector median 2.1); FY2022 fatality. "
         f"ISO 45001 at only 60% of sites.",
         "Extend ISO 45001 to 100% of sites; LTIFR target ≤1.5 by FY2026; fatality root-cause review."),
        ("MEDIUM", "Living Wage",
         "Contractor and sub-contractor workforce; sub-tier suppliers in lower-income countries.",
         "Living wage benchmark analysis; inclusion in CoC; supplier capacity building."),
        ("MEDIUM", "Community Land & Livelihood Rights",
         "Taranto site — historical contamination; local health impacts. Bilbao water-stress basin affecting local communities.",
         "Community Impact Assessment; Taranto remediation roadmap; Bilbao water reduction plan."),
        ("MEDIUM", "Non-Discrimination & Equal Opportunity",
         f"Female workforce 22% (below 28% sector peer); pay equity not publicly disclosed.",
         "Pay gap analysis; gender parity targets; EU Pay Transparency Directive readiness."),
        ("LOWER", "Privacy & Data Rights",
         "Employee monitoring systems; contractor data in digital procurement platforms.",
         "GDPR data processing review; privacy impact assessments for new systems."),
    ]
    headers = ["Severity", "Issue", f"{D.COMPANY['short']} Risk Context", "Key Mitigation Action"]
    widths  = [2.0, 3.2, 5.4, 5.4]
    tbl     = _make_table(doc, headers, widths)
    sev_fills = {"HIGH": "FAD7D4", "MEDIUM": "FEF9E7", "LOWER": "D5F5E3"}
    sev_colors = {"HIGH": B["red"], "MEDIUM": B["amber"], "LOWER": B["green"]}
    for i, (sev, issue, context, action) in enumerate(salient_rows):
        fill = sev_fills.get(sev, B["band_bg"])
        row  = tbl.add_row()
        data = [sev, issue, context, action]
        cws  = widths
        for j, (val, w) in enumerate(zip(data, cws)):
            cell = row.cells[j]
            cell.width = Cm(w)
            _cell_shade(cell, fill if j == 0 else (B["band_bg"] if i % 2 == 0 else "FFFFFF"))
            _cell_borders(cell, B["rule"])
            if j == 0:
                _cell_text(cell, val, bold=True, color=sev_colors.get(sev, B["slate"]), center=True)
            else:
                _cell_text(cell, val, bold=(j == 1))
    doc.add_paragraph()

    # ── Section 5: HRDD Process ─────────────────────────────────────────────
    _add_heading(doc, "5.  Human Rights Due Diligence Process", level=1,
                 page_break_before=True)
    _add_body(doc, (
        f"{D.COMPANY['short']}'s HRDD process follows the five-step cycle mandated by the UNGPs and "
        "operationalised under the CSDDD. The process applies to own operations and "
        "extends to the value chain on a risk-proportionate basis."
    ), space_after=6)

    # Process flow diagram
    flow_img = _chart_hrdd_process(tmpdir)
    _insert_image(doc, flow_img, width_cm=15,
                  caption="Figure 1: HRDD Five-Step Cycle (UNGP Pillar II  ·  CSDDD Art. 5–11)")

    hrdd_rows = [
        ("1", "Identify & Map",
         "Conduct value-chain mapping to identify all operations, suppliers, and business partners with potential human rights exposure.",
         "Annual; triggered by new sourcing geographies or supplier relationships.",
         "CSO + Procurement",
         "Value-chain risk map; updated supplier register."),
        ("2", "Assess Impacts",
         "Perform Human Rights Impact Assessment (HRIA) using IFC Performance Standards and OECD DD Guidance. Prioritise DRC minerals and Brazil iron-ore sourcing. Score severity × likelihood.",
         "Annual baseline; real-time for material changes.",
         "CSO + External HRDD Specialist",
         "HRIA report; salient issues register."),
        ("3", "Integrate & Act",
         "Embed HRDD findings into procurement decisions, contracts (CoC clauses), capital allocation, and supplier corrective action plans (CAPs). Apply leverage to influence suppliers.",
         "Continuous; CAPs within 90 days of audit finding.",
         "Procurement + Legal + Operations",
         "CAP tracker; procurement policy update; contracts amended."),
        ("4", "Track Progress",
         "Monitor effectiveness of actions through KPIs (audit coverage, CAP closure rate, grievance resolution rate). Internal audit and third-party verification for high-risk suppliers.",
         "Quarterly KPI review; annual third-party verification.",
         "CSO + Internal Audit",
         "KPI dashboard; audit reports."),
        ("5", "Communicate & Report",
         "Disclose HRDD process, salient issues, actions taken, and outcomes in ESRS S2/S3 disclosures, Annual Report, and Sustainability Report. Engage stakeholders on findings.",
         "Annual (ESRS report); ad hoc for material incidents.",
         "CSO + CFO + Communications",
         "ESRS S2/S3 disclosures; Sustainability Report chapter."),
    ]
    headers = ["Step", "Phase", "Activities", "Frequency", "Lead", "Output"]
    widths  = [1.0, 2.5, 4.5, 2.5, 2.5, 3.0]
    tbl     = _make_table(doc, headers, widths)
    for i, row_data in enumerate(hrdd_rows):
        _add_table_row(tbl, list(row_data), widths, row_idx=i)
    doc.add_paragraph()

    # ── Section 6: Supplier Code of Conduct ─────────────────────────────────
    _add_heading(doc, "6.  Supplier Code of Conduct Requirements", level=1,
                 page_break_before=True)
    _add_body(doc, (
        "All Tier-1 suppliers and high-risk Tier-2 suppliers must acknowledge and comply with "
        f"{D.COMPANY['short']}'s Supplier Code of Conduct (CoC) as a condition of doing business. "
        "The following minimum standards are non-negotiable:"
    ), space_after=6)

    coc_rows = [
        ("Child Labour (ILO C138/C182)",
         "PROHIBITED — No employment of persons under 15 (or higher local minimum age). No hazardous work for persons under 18.",
         "Contractual clause + SAQ declaration + on-site audit for high-risk suppliers."),
        ("Forced & Bonded Labour (ILO C29/C105)",
         "PROHIBITED — No use of forced, compulsory, bonded, trafficked, or prison labour. Workers retain freedom to leave employment.",
         "Contractual clause; worker interviews during audits; recruitment agency monitoring."),
        ("Freedom of Association & Collective Bargaining (ILO C87/C98)",
         "REQUIRED — Workers' right to form and join unions; right to bargain collectively. Non-interference and non-retaliation guaranteed.",
         "SAQ question; audit interview; trade union engagement in-country where possible."),
        ("Non-Discrimination (ILO C100/C111)",
         "REQUIRED — Equal treatment and opportunity. No discrimination on grounds of gender, race, religion, age, disability, sexual orientation, or trade union membership.",
         "SAQ; HR policy review; audit."),
        ("Living Wage",
         "REQUIRED — Workers to receive wages meeting or exceeding the applicable legal minimum wage; aspiration towards living wage benchmarks (e.g., WageIndicator, Anker Methodology).",
         "Payroll sampling during audits; Living Wage gap analysis for priority suppliers."),
        ("Occupational Health & Safety (ILO C155/ILO-OSH 2001)",
         "REQUIRED — Safe and healthy working conditions. Risk assessments, PPE provision, emergency procedures, incident reporting. Aspire to ISO 45001.",
         "Audit assessment; incident rate monitoring; ISO 45001 evidence for large suppliers."),
        ("Working Hours (ILO C1/C30)",
         "REQUIRED — Working hours compliant with national law and ILO standards. Overtime voluntary and compensated; maximum 60 hours/week including overtime.",
         "Time-record review during audits."),
        ("Environmental Baseline",
         "REQUIRED — Comply with applicable environmental laws; manage hazardous materials safely; no illegal disposal of waste.",
         "SAQ; environmental section of social audits."),
        ("Anti-Bribery & Corruption",
         "REQUIRED — Zero tolerance for bribery, corruption, facilitation payments. Comply with local law and FCPA/UK Bribery Act principles.",
         "Certification; reference-check for high-risk geographies."),
        ("Grievance Mechanism",
         f"REQUIRED — Suppliers to provide workers with access to an effective grievance mechanism free from retaliation. Notify {D.COMPANY['short']} of any material grievance relating to our supply.",
         "SAQ; verification during audit; incident reporting requirement in contracts."),
    ]
    headers = ["CoC Requirement", "Standard / Prohibition", "Verification Method"]
    widths  = [3.5, 6.5, 6.0]
    tbl     = _make_table(doc, headers, widths)
    for i, row_data in enumerate(coc_rows):
        _add_table_row(tbl, list(row_data), widths, row_idx=i)
    doc.add_paragraph()

    _add_heading(doc, "6.1  Supplier Engagement & Escalation", level=2)
    escalation_rows = [
        ("Tier 1 — All suppliers",    "Self-Assessment Questionnaire (SAQ) completion; annual recertification."),
        ("Tier 1 — High-risk",        "Third-party social audit (SMETA 4-Pillar or equivalent) within 24 months of onboarding."),
        ("Tier 2 — High-risk (DRC / Brazil)", "Enhanced HRDD: site visits, worker interviews, conflict-minerals traceability (OECD DD Guidance for Minerals)."),
        ("Breach — Remedial",         f"Corrective Action Plan (CAP) required within 30 days; {D.COMPANY['short']} support provided; 90-day resolution target."),
        ("Breach — Termination",      "Persistent or severe breach (e.g., forced/child labour confirmed) → suspension pending investigation → termination with 30-day notice."),
    ]
    headers = ["Tier / Scenario", "Engagement Action"]
    widths  = [5.0, 11.0]
    tbl2    = _make_table(doc, headers, widths)
    for i, row_data in enumerate(escalation_rows):
        _add_table_row(tbl2, list(row_data), widths, row_idx=i, bold_first=True)
    doc.add_paragraph()

    # ── Section 7: Grievance Mechanism ──────────────────────────────────────
    _add_heading(doc, "7.  Grievance Mechanism", level=1,
                 page_break_before=True)
    _add_body(doc, (
        f"{D.COMPANY['short']} is committed to providing access to effective, legitimate, and impartial "
        "grievance mechanisms for all those potentially affected by our human rights impacts, "
        "in line with UNGP Principle 29–31."
    ), space_after=6)

    _add_heading(doc, "7.1  Available Channels", level=2)
    channel_rows = [
        ("Dedicated Grievance Hotline",   "24/7 phone line (Italian + English + Portuguese + Swahili)",  "Confidential; anonymous option available", "All stakeholders"),
        ("Online Reporting Portal",        "Web-based form at [company.intranet/speak-up]",               "Anonymous; encrypted; mobile-friendly",    "Employees, contractors, suppliers"),
        ("Email Channel",                  "humanrights@verdasteelco.it",                                 "Confidential; named or anonymous",         "All stakeholders"),
        ("Local Site Contact",             "Designated HR/CSO representative at each site",               "In-person or written",                     "Employees, local community"),
        ("Community Liaison Officer",      "Appointed for Taranto and Bilbao given community exposure",   "In-person; local language",                "Communities"),
        ("Trade Union Representatives",    "Formalised union grievance procedure under Italian/Spanish collective agreements", "Formal; documented", "Employees"),
        ("External — ANAC / Authorities",  "Italy's ANAC (where governance concerns overlap) or competent national authority", "Independent external channel", "Any person"),
    ]
    headers = ["Channel", "Access Method", "Confidentiality", "Who Can Use"]
    widths  = [3.5, 4.5, 4.0, 4.0]
    tbl     = _make_table(doc, headers, widths)
    for i, row_data in enumerate(channel_rows):
        _add_table_row(tbl, list(row_data), widths, row_idx=i)
    doc.add_paragraph()

    _add_heading(doc, "7.2  Timelines & Process Standards", level=2)
    timeline_rows = [
        ("Acknowledgement",             "Within 7 calendar days of receipt",       "Grievance Mechanism Manager"),
        ("Initial Assessment / Triage", "Within 15 calendar days",                 "CSO / HR  /  Legal"),
        ("Investigation Launch",        "Within 20 calendar days of receipt",       "Independent Investigator (internal or external)"),
        ("Progress Update to Reporter", "Within 45 calendar days",                  "Grievance Mechanism Manager"),
        ("Final Feedback",              "Within 3 months (90 calendar days)",       "CSO"),
        ("Remediation Plan",            "Within 30 days of finding confirmation",   "Operations / Procurement / HR"),
        ("Closure & Learning",          "Post-resolution; lessons fed into HRDD cycle", "CSO"),
    ]
    headers = ["Stage", "Timeline", "Responsible Party"]
    widths  = [5.0, 5.0, 6.0]
    tbl2    = _make_table(doc, headers, widths)
    for i, row_data in enumerate(timeline_rows):
        _add_table_row(tbl2, list(row_data), widths, row_idx=i, bold_first=True)
    doc.add_paragraph()

    _add_heading(doc, "7.3  Non-Retaliation Commitment", level=2)
    _add_callout(
        doc,
        text=(
            f"{D.COMPANY['short']} strictly prohibits any form of retaliation — dismissal, demotion, "
            "harassment, disadvantage, or discriminatory treatment — against any person who raises "
            "a concern in good faith. This protection extends to employees, contractors, suppliers, "
            "community members, and third parties. Violations of this principle will be treated as "
            "serious misconduct subject to disciplinary action up to and including termination."
        ),
        label="NON-RETALIATION — ABSOLUTE COMMITMENT",
        fill="E8F1ED",
        border_color=B["accent"],
        label_color=B["accent"],
    )

    # ── Section 8: Roles & Responsibilities ────────────────────────────────
    _add_heading(doc, "8.  Roles & Responsibilities", level=1,
                 page_break_before=True)
    _add_body(doc, (
        "The following RACI table defines accountability for each key element of this Policy. "
        "R = Responsible  ·  A = Accountable  ·  C = Consulted  ·  I = Informed"
    ), space_after=6)

    raci_rows = [
        ("Approve this Policy",                    "A", "—",  "C",  "C",  "I"),
        ("HRDD programme oversight",               "C", "R/A","I",  "C",  "I"),
        ("Value-chain risk mapping",               "I", "A",  "R",  "C",  "I"),
        ("Supplier CoC implementation & audits",   "I", "A",  "R",  "C",  "I"),
        ("Grievance mechanism — operations",       "I", "A",  "C",  "R",  "I"),
        ("Legal / regulatory compliance monitoring","C", "C",  "I",  "A/R","I"),
        ("ESRS S2/S3 disclosures",                 "I", "A/R","C",  "C",  "I"),
        ("Employee training delivery",             "I", "A",  "C",  "C",  "R"),
        ("Board reporting (annual)",               "A", "R",  "C",  "I",  "I"),
        ("External stakeholder engagement",        "A", "R",  "C",  "I",  "I"),
    ]
    headers = ["Activity", "Board / ESG Committee", "CSO", "Procurement", "General Counsel", "HR / Ops"]
    widths  = [5.0, 2.5, 2.0, 2.5, 2.5, 1.5]
    tbl     = _make_table(doc, headers, widths)
    for i, row_data in enumerate(raci_rows):
        _add_table_row(tbl, list(row_data), widths, row_idx=i, bold_first=True)
    doc.add_paragraph()

    # ── Section 9: Training & Communication ────────────────────────────────
    _add_heading(doc, "9.  Training & Communication", level=1)
    training_rows = [
        ("All employees",          "Human Rights Awareness (online)",        "Annual (30 min)",       "Mandatory; completion tracked in LMS"),
        ("Procurement team",       "Supplier HRDD & Responsible Sourcing",   "Annual (2 hrs)",        "Mandatory; includes DRC/Brazil module"),
        ("Site managers / H&S leads", "Human Rights in Operations",          "Biennial (half-day)",   "Mandatory; site-specific content"),
        ("Senior leadership / Board", "HRDD Oversight & CSDDD Liability",    "Annual (1 hr briefing)","Board ESG Committee; CSO-led"),
        ("New joiners",            "Human Rights & Code of Conduct induction","Within 60 days",       "Mandatory for all new employees"),
        ("Suppliers (high-risk)",  "Supplier capacity building workshop",     "Annual; or post-audit", "Offered in local language; Portuguese + French + Swahili"),
    ]
    headers = ["Audience", "Programme", "Frequency", "Notes"]
    widths  = [3.5, 4.5, 3.0, 5.0]
    tbl     = _make_table(doc, headers, widths)
    for i, row_data in enumerate(training_rows):
        _add_table_row(tbl, list(row_data), widths, row_idx=i)
    doc.add_paragraph()

    # ── Section 10: Monitoring, KPIs & Review ──────────────────────────────
    _add_heading(doc, "10.  Monitoring, KPIs & Annual Review", level=1,
                 page_break_before=True)
    _add_body(doc, (
        "Performance against this Policy will be tracked through the following KPIs, "
        "reported quarterly to the CSO and annually to the Board ESG Committee and in "
        "ESRS S2/S3 disclosures."
    ), space_after=6)

    kpi_rows = [
        ("Supplier CoC coverage — Tier 1",       "% of Tier-1 spend covered by signed CoC", "0% (baseline)", "100% by end of Year 1"),
        ("Supplier SAQ completion rate",          "% of Tier-1 suppliers completing SAQ",     "0% (baseline)", "80% by end of Year 1; 100% Year 2"),
        ("High-risk supplier audit coverage",     "% of high-risk suppliers with audit/HRIA in last 24m", "0%", "60% by Year 2; 100% by Year 3"),
        ("Grievance mechanism — accessibility",  "% of employees aware of channel (survey)",  "Not measured", "≥80% by end of Year 1"),
        ("Grievance resolution rate",             "% of grievances closed within 90-day SLA", "Not tracked",  "≥90% by end of Year 1"),
        ("LTIFR (own workforce)",                 "Lost-Time Injury Frequency Rate (per M hrs)", "2.8",       "≤2.0 Year 2; ≤1.5 Year 3"),
        ("Fatalities",                            "Number of work-related fatalities",          "0.33/yr avg", "Zero target; root-cause actions"),
        ("ISO 45001 site coverage",              "% of sites certified ISO 45001",             "60%",         "80% Year 2; 100% Year 3"),
        ("HRDD training completion",             "% of target audience completing annual training", "0%",    "≥90% by end of Year 1"),
        ("ESRS S2/S3 disclosure completeness",  "% of required ESRS S2/S3 datapoints disclosed", "Not assessed", "Gap-close report within 6 months"),
    ]
    headers = ["KPI", "Metric Definition", "Baseline (FY2023)", "Target"]
    widths  = [4.5, 4.5, 2.5, 4.5]
    tbl     = _make_table(doc, headers, widths)
    for i, row_data in enumerate(kpi_rows):
        _add_table_row(tbl, list(row_data), widths, row_idx=i, bold_first=True)
    doc.add_paragraph()

    _add_heading(doc, "10.1  Policy Review Cycle", level=2)
    _add_body(doc, (
        "This Policy will be reviewed annually by the CSO and submitted to the Board ESG Committee "
        "for approval. An expedited review will be triggered by: (i) material change in operations "
        "or sourcing geographies; (ii) significant regulatory development (including CSDDD "
        "transposition); (iii) material human rights incident; or (iv) findings from external audit "
        "or stakeholder engagement indicating Policy gaps."
    ), space_after=6)

    # ── Appendix: 12-Month Implementation Roadmap ────────────────────────────
    _add_heading(doc, "Appendix A — 12-Month Implementation Roadmap",
                 level=1, page_break_before=True,
                 color_key="accent")
    _add_body(doc, (
        "The following roadmap translates this Policy into concrete actions. "
        "Total estimated investment: €250,000–€500,000 (per ESGIntel analysis, SOC-01 action plan)."
    ), space_after=6)

    roadmap_rows = [
        ("Month 1–2",   "Board approval of this Policy; appoint Human Rights Lead within CSO team; "
                        "launch internal communications campaign.",
         "CSO + Board", "High"),
        ("Month 1–3",   "Conduct value-chain risk mapping; identify top 20 high-risk suppliers "
                        "(Brazil iron-ore, DRC mineral chain).",
         "CSO + Procurement", "Critical"),
        ("Month 2–4",   "Develop and deploy Supplier Self-Assessment Questionnaire (SAQ); "
                        "distribute to all Tier-1 suppliers with CoC.",
         "Procurement + Legal", "High"),
        ("Month 3–5",   "Establish Grievance Mechanism: select platform vendor; configure channels "
                        "(hotline, portal, email); go live; communicate to all stakeholders.",
         "CSO + HR + IT", "Critical"),
        ("Month 3–6",   "Commission HRIA for DRC mineral supply chain (external specialist); "
                        "deliver findings report.",
         "CSO + External Consultant", "Critical"),
        ("Month 4–6",   "Deliver Human Rights awareness training to all employees; "
                        "specialist training for Procurement team.",
         "HR + CSO", "High"),
        ("Month 5–8",   "Commission Tier-1 social audits for top 5 high-risk suppliers "
                        "(SMETA or equivalent); receive audit reports; issue CAPs where required.",
         "Procurement", "High"),
        ("Month 6–9",   "Extend ISO 45001 certification programme to Brescia and Genoa sites; "
                        "begin process for Taranto (most complex).",
         "Operations + H&S", "High"),
        ("Month 9–11",  "First quarterly KPI review with Board ESG Committee; "
                        "mid-year grievance mechanism effectiveness assessment.",
         "CSO + Board ESG Committee", "Medium"),
        ("Month 11–12", "ESRS S2/S3 gap analysis and disclosure drafting for FY2025 report; "
                        "annual Policy review; submit revised Policy for Board approval.",
         "CSO + CFO", "High"),
    ]
    headers = ["Timeline", "Action", "Owner", "Priority"]
    widths  = [2.5, 8.5, 4.0, 1.0]  # sum=16
    tbl     = _make_table(doc, headers, widths)
    p_fills = {"Critical": "FAD7D4", "High": "FEF9E7", "Medium": "D5F5E3"}
    for i, (timeline, action, owner, priority) in enumerate(roadmap_rows):
        fill = p_fills.get(priority, B["band_bg"])
        row  = tbl.add_row()
        for j, (val, w) in enumerate(zip([timeline, action, owner, priority], widths)):
            cell = row.cells[j]
            cell.width = Cm(w)
            bg = fill if j == 3 else (B["band_bg"] if i % 2 == 0 else "FFFFFF")
            _cell_shade(cell, bg)
            _cell_borders(cell, B["rule"])
            _cell_text(cell, val, bold=(j == 3),
                       color=sev_colors.get(priority, B["slate"]) if j == 3 else B["slate"])
    doc.add_paragraph()

    _add_body(doc, (
        "Note: This document was prepared with AI assistance via the ESGIntel platform. "
        "It constitutes a substantive draft only. Legal review by qualified counsel "
        "familiar with Italian, Spanish, and EU law is required prior to adoption and publication."
    ), italic=True, color_key="muted", size=9, space_after=6)

    # ── Footer ──────────────────────────────────────────────────────────────
    _add_footer(doc, "Human Rights Policy")
    _add_page_numbers(doc)

    # ── Save ────────────────────────────────────────────────────────────────
    fname = os.path.join(out_dir, "Human Rights Policy Template.docx")
    doc.save(fname)
    return fname


# ════════════════════════════════════════════════════════════════════════════
# WHISTLEBLOWER / SPEAK-UP POLICY BUILDER
# ════════════════════════════════════════════════════════════════════════════

def _build_whistleblower(out_dir: str, tmpdir: str) -> str:
    C   = cd.COMPANY
    doc = _new_doc()

    # ── Title page ──────────────────────────────────────────────────────────
    _title_page(
        doc,
        policy_title="Whistleblower / Speak-Up Policy",
        subtitle=(
            "EU Directive 2019/1937  ·  Italy D.Lgs. 24/2023  ·  GDPR-Compliant"
        ),
        version="1.0",
        status="IMMEDIATE PRIORITY — NON-COMPLIANT SINCE DEC 2021",
    )
    doc.add_page_break()

    # ── Document control ────────────────────────────────────────────────────
    _add_heading(doc, "Document Control & Version History", level=1)
    _doc_control_table(
        doc,
        version="1.0  (Initial Draft)",
        owner="General Counsel",
        approved_by="Board of Directors / Administrative Body (Organo Amministrativo)",
        effective_date="[To be completed — TARGET: within 60 days of this draft]",
        review_date="Annually; upon material regulatory change",
        classification="Public",
    )

    # Compliance urgency callout
    risk_gov01 = next(r for r in cd.RISKS if r["id"] == "GOV-01")
    _add_callout(
        doc,
        text=(
            f"IMMEDIATE ACTION REQUIRED.  Risk ID GOV-01 — Score {risk_gov01['score']}/100 (High).  "
            f"{risk_gov01['summary']}  "
            f"EU Directive 2019/1937 required transposition by 17 December 2021; "
            f"Italy enacted D.Lgs. 24/2023 (effective 17 July 2023, mandatory for companies "
            f"with 50+ employees). {D.COMPANY['short']} ({C['employees']:,} employees) is an outlier "
            f"vs all sector peers and faces regulatory and reputational sanctions.  "
            f"Action: Deploy this Policy and a confidential reporting channel within 60 days."
        ),
        label="REGULATORY NON-COMPLIANCE  —  IMMEDIATE REMEDIATION REQUIRED",
        fill="FAD7D4",
        border_color=B["red"],
        label_color=B["red"],
    )

    doc.add_page_break()

    # ── Section 1: Purpose ──────────────────────────────────────────────────
    _add_heading(doc, "1.  Purpose", level=1)
    _add_body(doc, (
        f"{C['name']} ('{D.COMPANY['short']}' or 'the Company') is committed to the highest standards "
        f"of ethical conduct, legal compliance, and corporate integrity. This Whistleblower / "
        f"Speak-Up Policy ('Policy') establishes a clear, safe, and confidential framework "
        f"through which any person may report actual or suspected violations of law, "
        f"regulations, or Company policies without fear of retaliation."
    ), space_after=6)
    _add_body(doc, (
        "This Policy is adopted in compliance with:"
    ), space_after=3)
    fw_list = [
        "EU Directive 2019/1937 of 23 October 2019 on the protection of persons who report breaches of Union law ('Whistleblowing Directive');",
        "Italy: D.Lgs. 24 of 10 March 2023 ('Decreto Whistleblowing') — mandatory for organisations with 50+ employees (effective 17 July 2023);",
        "Italy: D.Lgs. 231/2001 — Organisational Model for offence prevention (Model 231);",
        "EU General Data Protection Regulation (GDPR) 2016/679 and Italian D.Lgs. 196/2003 (Privacy Code);",
        "Applicable Spanish labour law for operations at Bilbao.",
    ]
    for f in fw_list:
        _add_bullet(doc, f)
    doc.add_paragraph()

    # ── Section 2: Scope ────────────────────────────────────────────────────
    _add_heading(doc, "2.  Scope — Who Is Covered", level=1)
    _add_body(doc, (
        "This Policy protects all persons who report concerns in good faith, including:"
    ), space_after=4)

    scope_rows = [
        ("Employees", f"All permanent, fixed-term, part-time, and apprenticeship employees across all {D.COMPANY['short']} entities (Italy and Spain)."),
        ("Temporary & Agency Workers", f"All temporary, agency, and seconded staff performing work for or at {D.COMPANY['short']}."),
        ("Contractors & Sub-contractors", f"All individuals performing work under contract (including sole traders and freelancers) for {D.COMPANY['short']}."),
        ("Suppliers & Business Partners", f"Persons working within the supply chain who have become aware of breaches in the context of their business relationship with {D.COMPANY['short']}."),
        ("Former Employees", "Persons who raise concerns about matters that arose during their employment, including those already dismissed."),
        ("Job Applicants", "Candidates who become aware of potential violations during a recruitment process."),
        ("Shareholders & Board Members", "Directors, board members, shareholders, and members of supervisory bodies."),
        ("Civil Society / Third Parties", f"NGOs, journalists, and community members with credible information about violations connected to {D.COMPANY['short']}'s activities."),
    ]
    headers = ["Who", "Description"]
    widths  = [4.0, 12.0]
    tbl     = _make_table(doc, headers, widths)
    for i, row_data in enumerate(scope_rows):
        _add_table_row(tbl, list(row_data), widths, row_idx=i, bold_first=True)
    doc.add_paragraph()

    # ── Section 3: Reportable Concerns ──────────────────────────────────────
    _add_heading(doc, "3.  What Can Be Reported", level=1,
                 page_break_before=True)
    _add_body(doc, (
        "The following categories of concerns may be reported through the channels described "
        "in Section 4. This list is illustrative, not exhaustive."
    ), space_after=6)

    concern_rows = [
        ("Financial Crime",         "Fraud, embezzlement, misappropriation of Company assets, tax evasion, money laundering, market manipulation, false accounting."),
        ("Bribery & Corruption",    "Offering, soliciting, or accepting bribes; facilitation payments; corrupt dealing with public officials (Italy anti-corruption law; D.Lgs. 231/2001)."),
        ("Labour & Human Rights",   "Child or forced labour in operations or supply chain; discrimination; harassment; violation of collective bargaining rights; health and safety violations resulting in serious risk."),
        ("Environmental",           "Illegal discharge, pollution, improper waste disposal, falsification of environmental data, violations of ETS/CBAM obligations."),
        ("Data Protection & Privacy","Breaches of GDPR; unauthorised access to or disclosure of personal data; unlawful employee surveillance."),
        ("Competition Law",         "Price-fixing, market allocation, bid-rigging, abuse of dominant position (Art. 101/102 TFEU; Italian Competition Act)."),
        ("Product Safety",          "Falsification of product quality certifications; supply of non-conforming steel products; safety data falsification."),
        ("Public Procurement",      "Corruption or irregularity in public procurement processes."),
        ("Other Regulatory Violations", f"Any breach of applicable law, regulation, or {D.COMPANY['short']} Code of Ethics / Conduct."),
    ]
    headers = ["Category", "Examples"]
    widths  = [3.5, 12.5]
    tbl     = _make_table(doc, headers, widths)
    concern_fills = {
        "Financial Crime": "FAD7D4",
        "Bribery & Corruption": "FAD7D4",
        "Labour & Human Rights": "FEF9E7",
        "Environmental": "FEF9E7",
    }
    for i, (cat, examples) in enumerate(concern_rows):
        fill = concern_fills.get(cat, B["band_bg"] if i % 2 == 0 else "FFFFFF")
        row  = tbl.add_row()
        for j, (val, w) in enumerate(zip([cat, examples], widths)):
            cell = row.cells[j]
            cell.width = Cm(w)
            _cell_shade(cell, fill)
            _cell_borders(cell, B["rule"])
            _cell_text(cell, val, bold=(j == 0))
    doc.add_paragraph()

    _add_callout(
        doc,
        text=(
            "This Policy does NOT cover: personal employment disputes unrelated to compliance violations "
            "(e.g., salary grievances, performance management); interpersonal conflicts that do not involve "
            "legal or ethical violations. For personal employment matters, please use the Company's HR "
            "grievance procedure. Whistleblowers should report in good faith — knowingly false reports "
            "may be subject to disciplinary action."
        ),
        label="SCOPE CLARIFICATION",
        fill="E8F1ED",
        border_color=B["accent"],
        label_color=B["accent"],
    )

    # ── Section 4: Reporting Channels ───────────────────────────────────────
    _add_heading(doc, "4.  Reporting Channels", level=1,
                 page_break_before=True)
    _add_body(doc, (
        f"{D.COMPANY['short']} provides multiple independent, accessible, and confidential reporting "
        "channels. Reporters may use any channel; anonymous reporting is permitted and protected."
    ), space_after=6)

    channel_rows = [
        ("1", "Internal Reporting Portal",
         "[speak-up.verdasteelco.it] — secure, encrypted web-based platform",
         "24/7", "Yes", "Anonymous two-way dialogue with case manager"),
        ("2", "Confidential Hotline",
         "+39 800 XXX XXX (IT) / +34 900 XXX XXX (ES) — multilingual IVR",
         "24/7", "Yes", "Operated by independent third-party provider"),
        ("3", "Designated Email",
         "integrity@verdasteelco.it — encrypted mailbox",
         "Business hours (reply ≤7 days)", "Named", "Encrypted; access limited to Compliance Officer"),
        ("4", "Written Submission",
         "Sealed letter to: Chief Compliance Officer, [Registered Address]",
         "Postal", "Named or Anonymous", "Sent directly to Compliance Officer; unopened by others"),
        ("5", "In-Person (General Counsel)",
         "Direct confidential appointment with General Counsel or nominated deputy",
         "By appointment", "Named", "Reserved for complex matters requiring immediate discussion"),
        ("6", "External — ANAC",
         "Autorità Nazionale Anticorruzione (ANAC) — anac.it/whistleblowing",
         "As required", "Yes", "Available if internal channels are inadequate or compromised; mandatory for public sector related matters"),
    ]
    headers = ["#", "Channel", "Access Method", "Availability", "Anonymous?", "Notes"]
    widths  = [0.6, 2.8, 4.5, 2.5, 1.8, 3.8]
    tbl     = _make_table(doc, headers, widths)
    for i, row_data in enumerate(channel_rows):
        _add_table_row(tbl, list(row_data), widths, row_idx=i)
    doc.add_paragraph()

    _add_callout(
        doc,
        text=(
            "Reporter's choice of channel is absolute — you may use any channel listed above. "
            "You are not required to report internally before going to ANAC or another external "
            f"authority. {D.COMPANY['short']} strongly encourages internal reporting first as it often "
            "leads to faster resolution, but this is never a prerequisite."
        ),
        label="REPORTER'S RIGHTS",
        fill="E8F1ED",
        border_color=B["accent"],
        label_color=B["accent"],
    )

    # ── Section 5: Confidentiality & Data Protection ─────────────────────────
    _add_heading(doc, "5.  Confidentiality & Data Protection (GDPR)", level=1,
                 page_break_before=True)
    _add_body(doc, (
        f"{D.COMPANY['short']} treats all reports and the identity of reporters with the utmost "
        "confidentiality. Personal data processed in connection with whistleblowing reports "
        "is handled in strict compliance with the EU General Data Protection Regulation "
        "(GDPR) 2016/679 and Italian D.Lgs. 196/2003."
    ), space_after=6)

    conf_rows = [
        ("Identity Protection",
         "The identity of the reporter will not be disclosed without explicit written consent, "
         "except where required by law (e.g., criminal proceedings) or ordered by a court. "
         "Disclosure, if unavoidable, will be communicated to the reporter in advance."),
        ("Information Shared in the Report",
         "Content of the report will be disclosed only to those persons who are strictly "
         "necessary for the investigation. Recipients are bound by confidentiality obligations."),
        ("Anonymous Reports",
         "Anonymous reports are accepted and will be investigated on their merits. However, "
         "anonymous reporters may receive less feedback due to the inability to contact them."),
        ("Data Controller",
         f"{C['name']} acts as the Data Controller. "
         "The Compliance Officer / General Counsel is the designated DPA contact for "
         "whistleblowing data. A separate Privacy Notice is available at [link]."),
        ("Retention",
         "Personal data relating to a report is retained for a maximum of 5 years from the "
         "date of closure of the case, unless a longer period is required by applicable law or "
         "ongoing legal proceedings (see Section 9)."),
        ("Third-Party Platform",
         "Where an external third-party platform is used (e.g., for the hotline or portal), "
         "a Data Processing Agreement (DPA) compliant with GDPR Art. 28 will be in place. "
         "Data will not be transferred outside the EU/EEA without adequate safeguards."),
    ]
    headers = ["Topic", "Commitment"]
    widths  = [3.5, 12.5]
    tbl     = _make_table(doc, headers, widths)
    for i, row_data in enumerate(conf_rows):
        _add_table_row(tbl, list(row_data), widths, row_idx=i, bold_first=True)
    doc.add_paragraph()

    # ── Section 6: Non-Retaliation ──────────────────────────────────────────
    _add_heading(doc, "6.  Non-Retaliation Protections", level=1)
    _add_callout(
        doc,
        text=(
            f"{D.COMPANY['short']} ABSOLUTELY PROHIBITS any form of retaliation against any person "
            "who makes a report in good faith under this Policy, regardless of the outcome of "
            "the investigation. This prohibition applies to all persons within the scope of "
            "this Policy, including managers, supervisors, directors, and the Company itself."
        ),
        label="ZERO-TOLERANCE ANTI-RETALIATION",
        fill="FAD7D4",
        border_color=B["red"],
        label_color=B["red"],
    )
    _add_body(doc, "Prohibited retaliatory acts include but are not limited to:", space_after=3)
    retaliation_items = [
        "Dismissal, suspension, demotion, or denial of promotion;",
        "Pay reduction, removal of benefits, or change in working hours against the reporter's will;",
        "Negative performance appraisal or references;",
        "Harassment, intimidation, threats, or hostile work environment;",
        "Disciplinary action, blacklisting, or public exposure of identity;",
        "Cancellation or non-renewal of contract (contractors and suppliers);",
        "Any other detrimental treatment connected to the report.",
    ]
    for item in retaliation_items:
        _add_bullet(doc, item)
    _add_body(doc, (
        "Legal protections: In addition to this Policy, reporters benefit from protections under "
        "D.Lgs. 24/2023, including reversal of the burden of proof (the Company must demonstrate "
        "that any adverse treatment was NOT connected to the report), and the right to interim "
        "measures ordered by a court. Retaliation by any Company employee will be treated as "
        "serious misconduct and may result in disciplinary action including termination."
    ), space_after=6)

    # ── Section 7: Investigation Process & Timelines ────────────────────────
    _add_heading(doc, "7.  Investigation Process & Timelines", level=1,
                 page_break_before=True)
    _add_body(doc, (
        "All reports received through any channel will be assessed and investigated in a "
        "fair, impartial, and timely manner. The Compliance Officer manages the process and "
        "ensures independence."
    ), space_after=6)

    flow_img = _chart_investigation_flow(tmpdir)
    _insert_image(doc, flow_img, width_cm=15.5,
                  caption="Figure 1: Investigation Process Flow & Statutory Timelines (D.Lgs. 24/2023)")

    inv_rows = [
        ("1. Receipt",
         "Report received via any channel. Automated acknowledgement sent (portal/email) or verbal acknowledgement (hotline).",
         "Compliance Officer",
         "Immediate (auto-acknowledgement for digital channels)"),
        ("2. Formal Acknowledgement",
         "Written acknowledgement sent to reporter (unless anonymous and no contact method). Report logged in the case management system.",
         "Compliance Officer",
         "Within 7 calendar days of receipt (statutory requirement — D.Lgs. 24/2023 Art. 5)"),
        ("3. Initial Triage & Admissibility",
         "Assess: Is the report within scope? Is there sufficient information to proceed? Are there conflict-of-interest issues requiring escalation?",
         "Compliance Officer + General Counsel",
         "Within 15 calendar days of receipt"),
        ("4. Investigation Launch",
         "Designate investigator (internal or external depending on subject matter). Preserve evidence. Notify relevant stakeholders on a need-to-know basis only.",
         "General Counsel / External Investigator",
         "Within 20 calendar days of receipt"),
        ("5. Investigation Conduct",
         "Gather facts, conduct interviews, review documentation, consult specialists. Maintain confidentiality throughout. Subject of the report has right to respond (prior to final conclusions).",
         "Lead Investigator",
         "Duration proportionate to complexity; standard target 60 days"),
        ("6. Findings & Conclusions",
         "Produce investigation report; recommend remediation actions (disciplinary, legal, process improvements). Submit to General Counsel and CSO.",
         "General Counsel + Compliance Officer",
         "Within 90 calendar days of receipt (statutory feedback deadline — D.Lgs. 24/2023)"),
        ("7. Feedback to Reporter",
         "Inform reporter of outcome and follow-up actions taken (subject to legal constraints). Anonymous reporters will receive general status via portal if two-way messaging enabled.",
         "Compliance Officer",
         "Within 3 months (90 calendar days) of receipt — statutory deadline"),
        ("8. Remediation & Follow-Up",
         "Implement disciplinary, corrective, and/or legal actions. Update process controls to prevent recurrence. Board ESG Committee notified for material matters.",
         "General Counsel + HR + Operations",
         "Remediation plan within 30 days of final report"),
        ("9. Case Closure & Lessons Learned",
         "Close case in system. Annual aggregate reporting to Board (anonymised). Lessons incorporated into training and risk management.",
         "Compliance Officer + CSO",
         "Post-conclusion; annual aggregate report"),
    ]
    headers = ["Stage", "Description", "Responsible", "Timeline"]
    widths  = [2.0, 7.0, 3.5, 3.5]
    tbl     = _make_table(doc, headers, widths)
    for i, row_data in enumerate(inv_rows):
        _add_table_row(tbl, list(row_data), widths, row_idx=i, bold_first=True)
    doc.add_paragraph()

    # ── Section 8: Roles & Responsibilities (RACI) ──────────────────────────
    _add_heading(doc, "8.  Roles & Responsibilities", level=1,
                 page_break_before=True)
    _add_body(doc, (
        "R = Responsible  ·  A = Accountable  ·  C = Consulted  ·  I = Informed"
    ), space_after=6)

    raci_rows = [
        ("Approve and own this Policy",             "A",  "R",  "—",  "I",  "I"),
        ("Design & implement reporting channels",   "I",  "A/R","C",  "—",  "—"),
        ("Receive and triage reports",              "I",  "A/R","C",  "—",  "—"),
        ("Investigate reports",                     "I",  "A",  "R",  "C",  "—"),
        ("Maintain confidentiality & GDPR compliance","I","R",  "C",  "A",  "—"),
        ("Provide feedback to reporters",           "I",  "A/R","—",  "C",  "—"),
        ("Impose disciplinary measures",            "A",  "C",  "—",  "I",  "R"),
        ("Annual aggregate Board reporting",        "A",  "R",  "—",  "I",  "I"),
        ("Employee awareness & training",           "I",  "A",  "—",  "C",  "R"),
        ("ANAC registration (if applicable)",       "I",  "A/R","—",  "C",  "—"),
        ("Monitor for retaliation",                 "I",  "A/R","—",  "C",  "C"),
    ]
    headers = ["Activity", "Board / Admin Body", "General Counsel / Compliance", "External Investigator", "DPO / Legal", "HR"]
    widths  = [5.0, 2.5, 3.5, 2.5, 2.0, 1.0]  # sum=16.5 → fits within 16.5cm content area
    tbl     = _make_table(doc, headers, widths)
    for i, row_data in enumerate(raci_rows):
        _add_table_row(tbl, list(row_data), widths, row_idx=i, bold_first=True)
    doc.add_paragraph()

    # ── Section 9: Record-Keeping ────────────────────────────────────────────
    _add_heading(doc, "9.  Record-Keeping", level=1)
    _add_body(doc, (
        f"{D.COMPANY['short']} maintains a secure, confidential case management system for all "
        "whistleblowing reports. Access is restricted to the Compliance Officer, General Counsel, "
        "and, on a strictly need-to-know basis, designated investigators."
    ), space_after=6)

    rk_rows = [
        ("Case register",       "Unique reference number, date received, channel, category, status, outcome.",
         "Case management system", "5 years from case closure"),
        ("Investigation files", "All evidence, interview notes, correspondence, investigator reports.",
         "Secure document repository (access-controlled)", "5 years from case closure (minimum)"),
        ("Reporter communications", "Acknowledgements, feedback letters, status updates.",
         "Encrypted; case management system", "5 years"),
        ("Training records",    "Completion records for all awareness and investigator training.",
         "HR Learning Management System (LMS)", "Duration of employment + 3 years"),
        ("Annual aggregate report", "Anonymised summary for Board (number of reports, categories, outcomes).",
         "Board minutes / Board pack", "Permanently (board records)"),
    ]
    headers = ["Record Type", "Content", "Storage", "Retention Period"]
    widths  = [3.5, 5.5, 4.0, 3.0]
    tbl     = _make_table(doc, headers, widths)
    for i, row_data in enumerate(rk_rows):
        _add_table_row(tbl, list(row_data), widths, row_idx=i, bold_first=True)
    doc.add_paragraph()

    # ── Section 10: Awareness & Training ────────────────────────────────────
    _add_heading(doc, "10.  Awareness & Training", level=1)
    _add_body(doc, (
        "Effective whistleblowing protection depends on everyone understanding their rights "
        "and obligations. The following programme ensures comprehensive awareness:"
    ), space_after=6)

    training_rows = [
        ("All employees",              "Speak-Up Policy & Channel Awareness",         "At induction + annually",    "Mandatory; completion tracked in LMS"),
        ("Managers & Supervisors",     "Anti-Retaliation + Manager Obligations",       "Annual (1 hr)",              "Mandatory; managers are a key risk point"),
        ("Compliance Officer / Legal", "Investigation techniques; D.Lgs. 24/2023 updates", "At appointment + as required", "Specialist; may use external provider"),
        ("Board / Senior Leadership",  "Governance obligations; aggregate report review", "Annual (30 min briefing)", "Presented by General Counsel"),
        ("Suppliers & Contractors",    "Supplier Speak-Up channels & rights",          "Included in CoC onboarding",  "Written notice; supported by Human Rights Policy"),
    ]
    headers = ["Audience", "Training Content", "Frequency", "Notes"]
    widths  = [3.5, 5.0, 3.0, 4.5]
    tbl     = _make_table(doc, headers, widths)
    for i, row_data in enumerate(training_rows):
        _add_table_row(tbl, list(row_data), widths, row_idx=i)
    doc.add_paragraph()

    _add_body(doc, (
        "Communication channels: This Policy and channel details will be displayed on the Company "
        "intranet, included in the employee handbook, posted on physical notice boards at each site "
        "(Taranto, Brescia, Genoa, Bilbao), and included in supplier onboarding packages. "
        "Supplier-facing materials will be made available in Italian, English, Portuguese, and "
        "Spanish."
    ), space_after=6)

    # ── Appendix: 60-Day Implementation Checklist ───────────────────────────
    _add_heading(doc, "Appendix A — 60-Day Implementation Checklist",
                 level=1, page_break_before=True,
                 color_key="accent")
    _add_callout(
        doc,
        text=(
            "Target: Full compliance with D.Lgs. 24/2023 within 60 days of Board approval "
            "of this Policy. Estimated cost of implementation: €15,000–€50,000 (platform "
            "licence + legal review + training). This is a low-cost, high-impact remediation "
            "that eliminates an outlier governance risk."
        ),
        label="60-DAY TARGET",
        fill="E8F1ED",
        border_color=B["accent"],
        label_color=B["accent"],
    )

    checklist_rows = [
        ("Day 1–5",   "Board or delegated body formally approves this Policy; appoints General Counsel as Policy Owner.",
         "Board / General Counsel", "Board resolution / minutes"),
        ("Day 1–10",  "Legal review of Policy text by external counsel familiar with D.Lgs. 24/2023.",
         "General Counsel + External Counsel", "Final Policy version"),
        ("Day 5–15",  "Select and contract internal reporting platform provider (e.g., EQS, NAVEX, or equivalent GDPR-compliant solution). Execute DPA.",
         "General Counsel + IT + Procurement", "Signed vendor agreement + DPA"),
        ("Day 10–20", "Configure reporting portal and hotline (multilingual: Italian, English, Spanish). Test user flows (anonymous and named).",
         "IT + Compliance Officer", "Platform go-live; test report completed"),
        ("Day 15–25", "Train Compliance Officer and designated investigators; ANAC registration (if applicable for public-sector related scope).",
         "General Counsel + External Provider", "Training completion; ANAC record"),
        ("Day 15–30", "Appoint Data Protection Officer liaison; complete GDPR Record of Processing Activities (RoPA) entry for whistleblowing system.",
         "DPO / Legal", "RoPA updated"),
        ("Day 20–35", "Communicate Policy and channels to all employees: intranet, all-staff email, notice boards (all 4 sites).",
         "HR + Communications", "Communication sent; evidence retained"),
        ("Day 25–40", "Mandatory Speak-Up awareness training for all employees (online module); mandatory anti-retaliation training for all managers.",
         "HR + L&D", "≥80% completion; LMS records"),
        ("Day 30–45", "Notify all contractors, suppliers, and business partners of Policy and available channels; include in supplier CoC communications.",
         "Procurement", "Notification sent; acknowledgements logged"),
        ("Day 40–55", "Internal test of full investigation workflow (simulated anonymous report); document lessons.",
         "Compliance Officer + IT", "Test report and lessons document"),
        ("Day 55–60", "Board confirmation of implementation completion; include compliance status in next Board pack and public ESG reporting.",
         "General Counsel + CSO", "Board confirmation; ESG update"),
    ]
    headers = ["Timeline", "Action", "Owner", "Deliverable / Evidence"]
    widths  = [2.0, 7.5, 3.5, 3.0]
    tbl     = _make_table(doc, headers, widths)
    for i, (timeline, action, owner, deliverable) in enumerate(checklist_rows):
        _add_table_row(tbl, [timeline, action, owner, deliverable], widths, row_idx=i, bold_first=True)
    doc.add_paragraph()

    # Legal disclaimer
    _add_body(doc, (
        "Note: This document was prepared with AI assistance via the ESGIntel platform. "
        "It constitutes a substantive draft only. Review by qualified Italian and EU legal counsel "
        "is required prior to adoption, given the specific requirements of D.Lgs. 24/2023 "
        f"and the operational specifics of {D.COMPANY['short']}."
    ), italic=True, color_key="muted", size=9, space_after=6)

    # ── Footer ──────────────────────────────────────────────────────────────
    _add_footer(doc, "Whistleblower / Speak-Up Policy")
    _add_page_numbers(doc)

    # ── Save ────────────────────────────────────────────────────────────────
    fname = os.path.join(out_dir, "Whistleblower Policy Template.docx")
    doc.save(fname)
    return fname


# ════════════════════════════════════════════════════════════════════════════
# PUBLIC ENTRY POINT
# ════════════════════════════════════════════════════════════════════════════

def build_policy(key: str, out_dir: str) -> str:
    """
    Generate a policy document.

    Parameters
    ----------
    key     : 'human_rights'  — Human Rights Policy
              'whistleblower' — Whistleblower / Speak-Up Policy
    out_dir : Directory where the .docx will be written (must exist).

    Returns
    -------
    Absolute path to the generated .docx file.
    """
    out_dir = os.path.abspath(out_dir)
    os.makedirs(out_dir, exist_ok=True)

    tmpdir = tempfile.mkdtemp(prefix="esgpol_")
    try:
        if key == "human_rights":
            return _build_human_rights(out_dir, tmpdir)
        elif key == "whistleblower":
            return _build_whistleblower(out_dir, tmpdir)
        else:
            raise ValueError(f"Unknown policy key: {key!r}. Must be 'human_rights' or 'whistleblower'.")
    finally:
        # Clean up temporary chart images
        import shutil
        shutil.rmtree(tmpdir, ignore_errors=True)
