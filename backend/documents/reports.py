"""
ESGIntel — Report Generator
===========================
Produces McKinsey-style ESG due-diligence reports (.docx primary, .pdf secondary)
for the demo company the assessed company S.p.A.

Public contract:
    build_report(report_type: str, fmt: str, out_dir: str) -> str
    report_type in {'full','exec','climate','engagement','riskregister'}
    fmt in {'docx','pdf'}
    returns absolute path to the generated file.
"""""

import os
import sys
import tempfile
import shutil
from pathlib import Path
from typing import Optional

# ── non-interactive matplotlib ──────────────────────────────────────────────
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np

# ── python-docx ─────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants

# ── reportlab ───────────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm, mm
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                 TableStyle, Image as RLImage, PageBreak,
                                 HRFlowable, KeepTogether)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY

# ── company data ─────────────────────────────────────────────────────────────
sys.path.insert(0, os.path.dirname(__file__))
import company_data as cd

# ────────────────────────────────────────────────────────────────────────────
# Colour helpers
# ────────────────────────────────────────────────────────────────────────────
B = cd.BRAND  # shorthand

def _rgb(key: str) -> RGBColor:
    h = B[key]
    r, g, b = cd.hex_rgb(h)
    return RGBColor(r, g, b)

def _hex(key: str) -> str:
    return B[key]

RAG_COLOR = {
    "Critical": _hex("red"),
    "High":     "E67E22",   # darker orange
    "Medium":   _hex("amber"),
    "Low":      _hex("green"),
}
RAG_FILL = {
    "Critical": "FAD7D4",
    "High":     "FDEBD0",
    "Medium":   "FEF9E7",
    "Low":      "D5F5E3",
}
STATUS_COLOR = {
    "red":   (_hex("red"),   "FAD7D4"),
    "amber": (_hex("amber"), "FEF9E7"),
    "green": (_hex("green"), "D5F5E3"),
}


# ════════════════════════════════════════════════════════════════════════════
# CHART GENERATORS  (matplotlib → temp PNG paths)
# ════════════════════════════════════════════════════════════════════════════

def _chart_pillar_scores(tmpdir: str) -> str:
    """Horizontal bar chart of E/S/G risk scores."""
    fig, ax = plt.subplots(figsize=(6, 2.8))
    pillars = list(cd.SCORES["pillars"].keys())
    scores  = [cd.SCORES["pillars"][p]["score"] for p in pillars]
    colors_b = [cd.hex_rgb01(RAG_COLOR[cd.SCORES["pillars"][p]["label"]]) for p in pillars]
    y = range(len(pillars))
    bars = ax.barh(y, scores, color=colors_b, height=0.5, edgecolor="white", linewidth=0.5)
    ax.set_yticks(list(y))
    ax.set_yticklabels(pillars, fontsize=10, color="#" + B["ink"])
    ax.set_xlim(0, 100)
    ax.set_xlabel("Risk Score (0–100)", fontsize=8, color="#" + B["muted"])
    ax.axvline(50, color="#" + B["rule"], linestyle="--", linewidth=0.8)
    for bar, score in zip(bars, scores):
        ax.text(bar.get_width() + 1.5, bar.get_y() + bar.get_height()/2,
                str(score), va="center", fontsize=9, color="#" + B["ink"], fontweight="bold")
    ax.set_title("Pillar Risk Scores", fontsize=10, color="#" + B["ink"], pad=8)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_visible(False)
    ax.tick_params(left=False)
    fig.tight_layout()
    path = os.path.join(tmpdir, "chart_pillars.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def _chart_financial_exposure(tmpdir: str) -> str:
    """Bar chart of financial exposures in €M."""
    fig, ax = plt.subplots(figsize=(6.5, 3.2))
    items  = cd.FINANCIAL_EXPOSURE
    labels = [i["driver"].replace("(", "\n(") for i in items]
    amounts = [i["amount_eur_m"] for i in items]
    pal = {
        "Critical":  cd.hex_rgb01(B["red"]),
        "High":      cd.hex_rgb01("E67E22"),
        "Strategic": cd.hex_rgb01(B["accent"]),
    }
    bar_colors = [pal.get(i["rating"], cd.hex_rgb01(B["amber"])) for i in items]
    x = range(len(labels))
    bars = ax.bar(x, amounts, color=bar_colors, width=0.6, edgecolor="white", linewidth=0.5)
    ax.set_xticks(list(x))
    ax.set_xticklabels(labels, fontsize=7.5, color="#" + B["ink"])
    ax.set_ylabel("€M", fontsize=9, color="#" + B["muted"])
    ax.set_title("Quantified Financial Exposure (€M)", fontsize=10, color="#" + B["ink"], pad=8)
    for bar, val in zip(bars, amounts):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 3,
                f"€{val}M", ha="center", fontsize=7.5, color="#" + B["ink"], fontweight="bold")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    legend_patches = [
        mpatches.Patch(color=pal["Critical"], label="Critical"),
        mpatches.Patch(color=pal["High"], label="High"),
        mpatches.Patch(color=pal["Strategic"], label="Strategic"),
    ]
    ax.legend(handles=legend_patches, fontsize=8, loc="upper left")
    fig.tight_layout()
    path = os.path.join(tmpdir, "chart_exposure.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def _chart_risk_severity(tmpdir: str) -> str:
    """Donut chart of risk severity distribution."""
    from collections import Counter
    counts = Counter(cd.severity(r["score"]) for r in cd.RISKS)
    labels = ["Critical", "High", "Medium", "Low"]
    sizes  = [counts.get(l, 0) for l in labels]
    clrs   = [cd.hex_rgb01(B["red"]), cd.hex_rgb01("E67E22"),
               cd.hex_rgb01(B["amber"]), cd.hex_rgb01(B["green"])]
    fig, ax = plt.subplots(figsize=(4, 3.5))
    wedges, texts, autotexts = ax.pie(
        sizes, labels=labels, colors=clrs, autopct="%1.0f%%",
        startangle=140, pctdistance=0.75,
        wedgeprops=dict(width=0.5, edgecolor="white", linewidth=1.5)
    )
    for t in texts:    t.set_fontsize(9);  t.set_color("#" + B["ink"])
    for t in autotexts: t.set_fontsize(8); t.set_color("white"); t.set_fontweight("bold")
    ax.set_title("Risk Severity Distribution", fontsize=10, color="#" + B["ink"], pad=8)
    fig.tight_layout()
    path = os.path.join(tmpdir, "chart_severity.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def _chart_climate_scenarios(tmpdir: str) -> str:
    """Grouped bar chart for NGFS climate scenario P&L."""
    fig, ax = plt.subplots(figsize=(6, 3))
    scenarios = [s["scenario"].split("(")[0].strip() for s in cd.CLIMATE_SCENARIOS]
    pnl       = [s["pnl_2030_eur_m"] for s in cd.CLIMATE_SCENARIOS]
    clrs      = [cd.hex_rgb01(B["green"]), cd.hex_rgb01(B["amber"]), cd.hex_rgb01(B["red"])]
    x = range(len(scenarios))
    bars = ax.bar(x, pnl, color=clrs, width=0.5, edgecolor="white")
    ax.set_xticks(list(x))
    ax.set_xticklabels(scenarios, fontsize=8.5, color="#" + B["ink"])
    ax.set_ylabel("P&L Impact (€M, 2030)", fontsize=8, color="#" + B["muted"])
    ax.axhline(0, color="#" + B["rule"], linewidth=0.8)
    ax.set_title("NGFS Climate Scenario P&L Impact (2030)", fontsize=10, color="#" + B["ink"], pad=8)
    for bar, val in zip(bars, pnl):
        ax.text(bar.get_x() + bar.get_width()/2, val - 3,
                f"€{val}M", ha="center", va="top", fontsize=8.5,
                color="white", fontweight="bold")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    path = os.path.join(tmpdir, "chart_climate.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def _chart_kpi_heatmap(tmpdir: str, group: Optional[str] = None) -> str:
    """Horizontal bar chart showing KPI peer percentiles."""
    kpis = [k for k in cd.KPIS if k["pctl"] is not None]
    if group:
        kpis = [k for k in kpis if k["group"].startswith(group)]
    if not kpis:
        return ""
    labels  = [k["metric"] for k in kpis]
    pctls   = [k["pctl"] for k in kpis]
    dirs    = [k["dir"] for k in kpis]
    bar_cols = []
    for pctl, d in zip(pctls, dirs):
        if d == "lower_better":
            risk = pctl
        else:
            risk = 100 - pctl
        if risk >= 75:   bar_cols.append(cd.hex_rgb01(B["red"]))
        elif risk >= 50: bar_cols.append(cd.hex_rgb01(B["amber"]))
        else:            bar_cols.append(cd.hex_rgb01(B["green"]))

    fig, ax = plt.subplots(figsize=(6.5, max(2.5, len(labels) * 0.45)))
    y = range(len(labels))
    ax.barh(list(y), pctls, color=bar_cols, height=0.55, edgecolor="white")
    ax.set_yticks(list(y))
    ax.set_yticklabels(labels, fontsize=8)
    ax.set_xlim(0, 100)
    ax.axvline(50, color="#" + B["rule"], linestyle="--", linewidth=0.8)
    ax.set_xlabel("Peer Percentile (0=best, 100=worst for lower_better)", fontsize=7)
    ax.set_title(f"KPI Peer Percentiles" + (f" — {group}" if group else ""), fontsize=9, color="#" + B["ink"])
    for pctl_v, yi in zip(pctls, y):
        ax.text(pctl_v + 1, yi, f"P{pctl_v}", va="center", fontsize=7.5)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_visible(False)
    ax.tick_params(left=False)
    fig.tight_layout()
    suffix = group.replace(" ", "_") if group else "all"
    path = os.path.join(tmpdir, f"chart_kpi_{suffix}.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


# ════════════════════════════════════════════════════════════════════════════
# DOCX HELPERS
# ════════════════════════════════════════════════════════════════════════════

def _new_doc() -> Document:
    doc = Document()
    # Page margins: A4 with 2.5cm margins
    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
    return doc


def _set_para_fmt(para, space_before=0, space_after=0, line_spacing=None, keep_together=False):
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if line_spacing:
        pf.line_spacing = Pt(line_spacing)
    if keep_together:
        pf.keep_together = True


def _add_heading(doc: Document, text: str, level: int = 1,
                 color_key: str = "ink", size: int = None,
                 space_before: int = 12, space_after: int = 6,
                 page_break_before: bool = False) -> None:
    para = doc.add_paragraph()
    if page_break_before:
        para.paragraph_format.page_break_before = True
    run = para.add_run(text)
    sizes = {1: 20, 2: 14, 3: 12, 4: 11}
    run.font.size = Pt(size or sizes.get(level, 11))
    run.font.bold = True
    run.font.color.rgb = _rgb(color_key)
    run.font.name = B["font_head"]
    _set_para_fmt(para, space_before=space_before, space_after=space_after)
    if level <= 2:
        # bottom border rule
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4" if level == 1 else "2")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), B["accent"])
        pBdr.append(bottom)
        pPr.append(pBdr)


def _add_body(doc: Document, text: str, bold: bool = False,
              italic: bool = False, color_key: str = "slate",
              size: int = 10, space_after: int = 4) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = _rgb(color_key)
    run.font.name = B["font_body"]
    _set_para_fmt(para, space_after=space_after)


def _add_caption(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = _rgb("muted")
    _set_para_fmt(para, space_before=2, space_after=8)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _cell_shade(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cell_borders(cell, color: str = None):
    color = color or B["rule"]
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _cell_text(cell, text: str, bold: bool = False, size: int = 9,
               color: str = None, center: bool = False) -> None:
    color = color or B["ink"]
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(str(text))
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.name  = B["font_body"]
    run.font.color.rgb = RGBColor(*cd.hex_rgb(color))
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(1)


def _add_table_header_row(table, headers: list, col_widths: list,
                           fill: str = None) -> None:
    fill = fill or B["ink"]
    row = table.rows[0]
    for i, (hdr, w) in enumerate(zip(headers, col_widths)):
        cell = row.cells[i]
        cell.width = Cm(w)
        _cell_shade(cell, fill)
        _cell_borders(cell, B["rule"])
        _cell_text(cell, hdr, bold=True, color="FFFFFF", center=True)


def _add_table_row(table, values: list, col_widths: list,
                   fill: str = "FFFFFF", row_idx: int = 0) -> None:
    row = table.add_row()
    bg  = fill if fill != "FFFFFF" else (B["band_bg"] if row_idx % 2 == 0 else "FFFFFF")
    for i, (val, w) in enumerate(zip(values, col_widths)):
        cell = row.cells[i]
        cell.width = Cm(w)
        _cell_shade(cell, bg)
        _cell_borders(cell, B["rule"])
        _cell_text(cell, str(val) if val is not None else "—")


def _insert_image(doc: Document, path: str, width_cm: float = 14,
                  caption: str = "") -> None:
    if not path or not os.path.exists(path):
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(path, width=Cm(width_cm))
    if caption:
        _add_caption(doc, caption)


def _add_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        run = para.add_run(f"{B['platform']} — CONFIDENTIAL   |   {cd.COMPANY['name']}   |   {cd.COMPANY['report_period']}")
        run.font.size = Pt(7.5)
        run.font.color.rgb = _rgb("muted")
        run.font.name = B["font_body"]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add hairline rule above footer
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        top = OxmlElement("w:top")
        top.set(qn("w:val"), "single")
        top.set(qn("w:sz"), "2")
        top.set(qn("w:space"), "4")
        top.set(qn("w:color"), B["rule"])
        pBdr.append(top)
        pPr.append(pBdr)


# ════════════════════════════════════════════════════════════════════════════
# SHARED DOCX SECTION BUILDERS
# ════════════════════════════════════════════════════════════════════════════

def _cover_page(doc: Document, title: str, subtitle: str = "") -> None:
    """Full-page branded cover."""
    # Top band spacer
    for _ in range(3):
        doc.add_paragraph()

    # Platform name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(B["platform"].upper())
    r.font.size  = Pt(11)
    r.font.bold  = True
    r.font.color.rgb = _rgb("accent")
    r.font.name  = B["font_body"]
    _set_para_fmt(p, space_after=4)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(B["tagline"])
    r2.font.size  = Pt(9)
    r2.font.color.rgb = _rgb("muted")
    r2.font.name  = B["font_body"]
    _set_para_fmt(p2, space_after=20)

    # Separator rule
    hr = doc.add_paragraph()
    hr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = hr._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "12")
    bot.set(qn("w:space"), "1");   bot.set(qn("w:color"), B["accent"])
    pBdr.append(bot); pPr.append(pBdr)
    _set_para_fmt(hr, space_after=20)

    # Company name
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(cd.COMPANY["name"])
    r3.font.size  = Pt(26)
    r3.font.bold  = True
    r3.font.color.rgb = _rgb("ink")
    r3.font.name  = B["font_head"]
    _set_para_fmt(p3, space_after=6)

    # Sector
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(cd.COMPANY["sector"])
    r4.font.size  = Pt(11)
    r4.font.color.rgb = _rgb("muted")
    r4.font.name  = B["font_body"]
    _set_para_fmt(p4, space_after=20)

    # Report title
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = p5.add_run(title)
    r5.font.size  = Pt(18)
    r5.font.bold  = True
    r5.font.color.rgb = _rgb("accent")
    r5.font.name  = B["font_head"]
    _set_para_fmt(p5, space_before=10, space_after=8)

    if subtitle:
        p6 = doc.add_paragraph()
        p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r6 = p6.add_run(subtitle)
        r6.font.size  = Pt(11)
        r6.font.color.rgb = _rgb("slate")
        r6.font.name  = B["font_body"]
        _set_para_fmt(p6, space_after=20)

    # Meta info
    for label, val in [
        ("Report Period", cd.COMPANY["report_period"]),
        ("As of",         cd.COMPANY["as_of"]),
        ("Domicile",      cd.COMPANY["domicile"]),
        ("Employees",     f"{cd.COMPANY['employees']:,}"),
        ("Revenue",       f"€{cd.COMPANY['revenue_eur_bn']}B"),
    ]:
        pm = doc.add_paragraph()
        pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rm1 = pm.add_run(f"{label}:  ")
        rm1.font.size = Pt(9); rm1.font.bold = True
        rm1.font.color.rgb = _rgb("ink"); rm1.font.name = B["font_body"]
        rm2 = pm.add_run(val)
        rm2.font.size = Pt(9)
        rm2.font.color.rgb = _rgb("slate"); rm2.font.name = B["font_body"]
        _set_para_fmt(pm, space_after=2)

    # Confidential banner
    for _ in range(3):
        doc.add_paragraph()
    pconf = doc.add_paragraph()
    pconf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rconf = pconf.add_run(B["confidential"].upper())
    rconf.font.size  = Pt(9)
    rconf.font.bold  = True
    rconf.font.color.rgb = _rgb("muted")
    rconf.font.name  = B["font_body"]

    doc.add_page_break()


def _scorecard_section(doc: Document) -> None:
    """At-a-glance ESG scorecard table."""
    _add_heading(doc, "ESG Risk Scorecard — At a Glance", level=1)

    s = cd.SCORES
    doc.add_paragraph()

    # Overall score box
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Overall ESG Risk Score:  {s['overall']} / 100   ({s['overall_label']})   P{s['overall_percentile']} vs Peers")
    r.font.size  = Pt(14)
    r.font.bold  = True
    r.font.color.rgb = _rgb("red")
    r.font.name  = B["font_head"]
    _set_para_fmt(p, space_before=4, space_after=10)

    # Pillar table
    headers = ["Pillar", "Risk Score", "Risk Level", "Peer Percentile", "Weight"]
    widths  = [3.5, 3.0, 3.0, 3.5, 2.5]
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_table_header_row(tbl, headers, widths)
    for row_idx, (pillar, data) in enumerate(s["pillars"].items()):
        row = tbl.add_row()
        vals = [pillar, f"{data['score']} / 100", data["label"],
                f"P{data['percentile']}", f"{int(data['weight']*100)}%"]
        for i, (v, w) in enumerate(zip(vals, widths)):
            cell = row.cells[i]
            cell.width = Cm(w)
            _cell_borders(cell, B["rule"])
            # colour the Risk Level cell
            if i == 2:
                _cell_shade(cell, RAG_FILL[data["label"]])
                _cell_text(cell, v, bold=True, color=RAG_COLOR[data["label"]])
            else:
                _cell_shade(cell, "F4F6F5" if row_idx % 2 == 0 else "FFFFFF")
                _cell_text(cell, v, bold=(i == 0))
    _set_para_fmt(doc.add_paragraph(), space_after=8)

    p2 = doc.add_paragraph()
    r2 = p2.add_run(f"Methodology: {s['formula']}")
    r2.font.size = Pt(8); r2.font.italic = True
    r2.font.color.rgb = _rgb("muted"); r2.font.name = B["font_body"]
    _set_para_fmt(p2, space_after=4)
    p3 = doc.add_paragraph()
    r3 = p3.add_run(f"Confidence: {s['confidence']}%  |  Peer universe: NACE C24 (n=47)  |  Period: {cd.COMPANY['report_period']}")
    r3.font.size = Pt(8); r3.font.italic = True
    r3.font.color.rgb = _rgb("muted"); r3.font.name = B["font_body"]
    _set_para_fmt(p3, space_after=10)


def _risk_register_table(doc: Document, risks=None, top_n=None) -> None:
    """Full or top-N risk register table."""
    data = risks or cd.RISKS
    if top_n:
        data = sorted(data, key=lambda r: r["score"], reverse=True)[:top_n]
    headers = ["ID", "Risk", "Sev.", "Score", "Horizon", "Financial Exposure", "Recommended Action"]
    widths  = [1.4, 3.6, 1.3, 1.1, 2.3, 3.2, 4.2]
    tbl = doc.add_table(rows=1, cols=7)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_table_header_row(tbl, headers, widths, fill=B["ink"])
    for idx, risk in enumerate(data):
        sev   = cd.severity(risk["score"])
        row   = tbl.add_row()
        bg    = "F4F6F5" if idx % 2 == 0 else "FFFFFF"
        vals  = [risk["id"], risk["title"], sev, str(risk["score"]),
                 risk["horizon"], risk["financial"], risk["action"]]
        for i, (v, w) in enumerate(zip(vals, widths)):
            cell = row.cells[i]
            cell.width = Cm(w)
            _cell_borders(cell, B["rule"])
            if i == 2:  # severity badge cell
                _cell_shade(cell, RAG_FILL[sev])
                _cell_text(cell, v, bold=True, color=RAG_COLOR[sev], center=True, size=8)
            elif i == 3:  # score
                _cell_shade(cell, bg)
                _cell_text(cell, v, bold=True, center=True)
            else:
                _cell_shade(cell, bg)
                _cell_text(cell, v, bold=(i == 0), size=8 if i in (5, 6) else 9)
    doc.add_paragraph()


def _financial_exposure_table(doc: Document) -> None:
    headers = ["Financial Driver", "Exposure (€M)", "Type", "Rating"]
    widths  = [5.5, 3.0, 3.5, 2.5]
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_table_header_row(tbl, headers, widths)
    for idx, item in enumerate(cd.FINANCIAL_EXPOSURE):
        sev = item["rating"]
        row = tbl.add_row()
        bg  = "F4F6F5" if idx % 2 == 0 else "FFFFFF"
        for i, (v, w) in enumerate(zip(
            [item["driver"], f"€{item['amount_eur_m']}M", item["type"], sev], widths)):
            cell = row.cells[i]
            cell.width = Cm(w)
            _cell_borders(cell, B["rule"])
            if i == 3:
                _cell_shade(cell, RAG_FILL.get(sev, "FFFFFF"))
                _cell_text(cell, v, bold=True, color=RAG_COLOR.get(sev, B["ink"]), center=True)
            else:
                _cell_shade(cell, bg)
                _cell_text(cell, v, bold=(i == 0))
    doc.add_paragraph()


def _exec_summary_text(doc: Document) -> None:
    _add_heading(doc, "Executive Summary", level=1)
    _add_body(doc,
        f"{cd.COMPANY['name']} ({cd.COMPANY['sector']}) presents a HIGH ESG risk profile with an "
        f"overall score of {cd.SCORES['overall']}/100 — bottom quintile (P{cd.SCORES['overall_percentile']}) "
        f"versus {cd.COMPANY['sector']} peers. Three ESG risk clusters drive the rating, each carrying "
        f"quantifiable financial consequences within the investment horizon.", size=10)
    _add_body(doc,
        "CLIMATE & CARBON (highest priority): With 842,000 tCO2e Scope 1 on a BF-BOF process route, "
        f"{D.COMPANY['short']} faces €68M/yr in net ETS costs by 2030 under the Orderly scenario, plus €18-28M "
        "additional CBAM exposure on non-EU exports. No internal carbon price, no SBTi commitment, and "
        "no DRI/EAF transition plan are in place. Asset stranding of the primary route is a credible "
        "risk by 2032-2035.", size=10, space_after=6)
    _add_body(doc,
        "COMPLIANCE & GOVERNANCE (immediate): Two immediate regulatory breaches — the EU whistleblower "
        "directive (non-compliant since December 2021) and CSRD/ESRS readiness (<18 months to first "
        "mandatory report) — represent remediable, low-cost risks that remain unaddressed. CSDDD "
        "exposure of up to €60M (5% turnover) materialises in 2027.", size=10, space_after=6)
    _add_body(doc,
        "MARKET ACCESS (revenue at risk): The absence of an SBTi target puts an estimated €456M "
        "(38% of revenue) of automotive-sector contracts at risk as OEM Scope 3 decarbonisation "
        "requirements tighten from 2025-2028. Green-bond and SLL market access is currently blocked, "
        "foreclosing €400-550M of potential lower-cost financing.", size=10, space_after=6)
    _add_body(doc,
        "OPPORTUNITIES: A credible SBTi commitment, ESRS E1 climate plan, and LTIFR improvement "
        "would unlock an EU Green Bond (€200-300M, 30-50bp saving), a Sustainability-Linked Loan "
        "(€200-250M, 25-35bp saving), and EU Innovation Fund grants for the DRI/EAF transition "
        "(€50-120M non-dilutive).", size=10, space_after=10)


def _kpi_section(doc: Document, pillar: str) -> None:
    """KPI table for a given pillar (Environmental/Social/Governance)."""
    group_map = {
        "Environmental": ["Climate (E1)", "Water (E3)", "Circular (E5)"],
        "Social":        ["Workforce (S1)"],
        "Governance":    ["Governance (G1)"],
    }
    groups = group_map.get(pillar, [])
    kpis   = [k for k in cd.KPIS if k["group"] in groups]
    if not kpis:
        return
    headers = ["Metric", "Group", f"{D.COMPANY['short']}", "Peer Median", "Peer Pctl", "Signal"]
    widths  = [4.0, 2.6, 2.8, 2.8, 2.0, 2.3]
    tbl = doc.add_table(rows=1, cols=6)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_table_header_row(tbl, headers, widths)
    for idx, kpi in enumerate(kpis):
        pctl = kpi["pctl"]
        d    = kpi["dir"]
        if pctl is None:
            signal, fill = "N/D", "F4F6F5"
        elif d == "lower_better":
            risk = pctl
        else:
            risk = 100 - pctl
        if pctl is not None:
            if risk >= 75:   signal, fill = "High Risk", RAG_FILL["High"]
            elif risk >= 50: signal, fill = "Medium",    RAG_FILL["Medium"]
            else:            signal, fill = "Low Risk",  RAG_FILL["Low"]
        bg = "F4F6F5" if idx % 2 == 0 else "FFFFFF"
        row = tbl.add_row()
        vals = [kpi["metric"], kpi["group"], kpi["value"],
                kpi["peer"] if kpi["peer"] else "—",
                f"P{pctl}" if pctl else "N/D", signal]
        for i, (v, w) in enumerate(zip(vals, widths)):
            cell = row.cells[i]
            cell.width = Cm(w)
            _cell_borders(cell, B["rule"])
            if i == 5:
                clr_map = {"High Risk": B["red"], "Medium": B["amber"],
                           "Low Risk": B["green"], "N/D": B["muted"]}
                _cell_shade(cell, fill if pctl else "F4F6F5")
                _cell_text(cell, v, bold=True, color=clr_map.get(v, B["ink"]), center=True, size=8)
            else:
                _cell_shade(cell, bg)
                _cell_text(cell, v, bold=(i == 0), size=9)
    doc.add_paragraph()


def _opportunities_table(doc: Document) -> None:
    headers = ["Lever", "Size", "Benefit", "Pre-condition"]
    widths  = [3.8, 2.5, 4.0, 4.5]
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_table_header_row(tbl, headers, widths, fill=B["accent"])
    for idx, opp in enumerate(cd.OPPORTUNITIES):
        row = tbl.add_row()
        bg  = "E8F1ED" if idx % 2 == 0 else "FFFFFF"
        for i, (v, w) in enumerate(zip(
            [opp["lever"], opp["size"], opp["benefit"], opp["precondition"]], widths)):
            cell = row.cells[i]
            cell.width = Cm(w)
            _cell_borders(cell, B["rule"])
            _cell_shade(cell, bg)
            _cell_text(cell, v, bold=(i == 0), size=9)
    doc.add_paragraph()


def _reg_timeline_table(doc: Document) -> None:
    headers = ["Year", "Regulatory Item", "Status"]
    widths  = [1.8, 10.0, 3.7]
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_table_header_row(tbl, headers, widths)
    for idx, item in enumerate(cd.REG_TIMELINE):
        row = tbl.add_row()
        bg  = "F4F6F5" if idx % 2 == 0 else "FFFFFF"
        st  = item["status"]
        st_color = B["red"] if "Non-comp" in st else (B["amber"] if "Imm" in st else B["ink"])
        for i, (v, w) in enumerate(zip([item["year"], item["item"], st], widths)):
            cell = row.cells[i]
            cell.width = Cm(w)
            _cell_borders(cell, B["rule"])
            _cell_shade(cell, bg)
            clr = st_color if i == 2 else B["ink"]
            _cell_text(cell, v, bold=(i == 0), color=clr, size=9)
    doc.add_paragraph()


def _engagement_asks_table(doc: Document) -> None:
    headers = ["Ask", "Owner", "Timeline", "Priority"]
    widths  = [7.0, 3.0, 2.5, 2.0]
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_table_header_row(tbl, headers, widths, fill=B["ink"])
    for idx, ask in enumerate(cd.ENGAGEMENT_ASKS):
        row = tbl.add_row()
        bg  = "F4F6F5" if idx % 2 == 0 else "FFFFFF"
        prio = ask["priority"]
        fill_map = {"Critical": RAG_FILL["Critical"], "High": RAG_FILL["High"],
                    "Immediate": RAG_FILL["Critical"]}
        prio_fill  = fill_map.get(prio, "FFFFFF")
        prio_color = RAG_COLOR.get(prio, B["ink"])
        for i, (v, w) in enumerate(zip(
            [ask["ask"], ask["owner"], ask["by"], prio], widths)):
            cell = row.cells[i]
            cell.width = Cm(w)
            _cell_borders(cell, B["rule"])
            if i == 3:
                _cell_shade(cell, prio_fill)
                _cell_text(cell, v, bold=True, color=prio_color, center=True, size=8)
            else:
                _cell_shade(cell, bg)
                _cell_text(cell, v, bold=(i == 0), size=9)
    doc.add_paragraph()


def _pillar_analysis(doc: Document, pillar: str, analysis_text: str,
                     tmpdir: str, chart_group: str = None) -> None:
    data = cd.SCORES["pillars"][pillar]
    score_label = f"{data['score']}/100 — {data['label']} Risk  (P{data['percentile']} vs peers)"
    _add_heading(doc, f"{pillar} Pillar", level=1, page_break_before=True)
    p = doc.add_paragraph()
    r = p.add_run(f"Risk Score: {score_label}")
    r.font.size = Pt(11); r.font.bold = True
    r.font.color.rgb = _rgb({"High": "red", "Medium": "amber", "Low": "green"}.get(data["label"], "ink"))
    r.font.name = B["font_head"]
    _set_para_fmt(p, space_after=6)
    _add_body(doc, analysis_text)
    _add_heading(doc, f"{pillar} KPIs vs Peer Universe", level=2)
    _kpi_section(doc, pillar)
    if chart_group:
        ch = _chart_kpi_heatmap(tmpdir, group=chart_group)
        if ch:
            _insert_image(doc, ch, width_cm=14,
                          caption=f"Figure: {pillar} KPI Peer Percentiles — Red=High Risk, Amber=Medium, Green=Low Risk")


# ════════════════════════════════════════════════════════════════════════════
# REPORT BUILDERS — DOCX
# ════════════════════════════════════════════════════════════════════════════

def _build_full_docx(out_path: str, tmpdir: str) -> None:
    doc = _new_doc()
    _add_footer(doc)

    # Cover
    _cover_page(doc, cd.REPORT_TYPES["full"],
                subtitle=f"ESG Due Diligence — {cd.COMPANY['report_period']}  |  {cd.COMPANY['as_of']}")

    # Scorecard
    _scorecard_section(doc)

    # Pillar chart
    ch_pillars = _chart_pillar_scores(tmpdir)
    _insert_image(doc, ch_pillars, width_cm=13,
                  caption="Figure 1: E/S/G Pillar Risk Scores — scores ≥55 = High Risk")

    # Executive Summary
    doc.add_page_break()
    _exec_summary_text(doc)

    # ── Environmental ──────────────────────────────────────────────────────
    _pillar_analysis(doc, "Environmental",
        (f"The Environmental pillar is the primary driver of ESG risk. {D.COMPANY['short']}'s 100% BF-BOF "
         "primary route generates 842,000 tCO2e Scope 1 annually — P89 vs sector peers — against a "
         "rising EU ETS price and a rapidly tapering free-allocation schedule (≈40% today → 10% by "
         "2030). No internal carbon price, no ETS hedging, and no SBTi commitment are in place. "
         "Water risk compounds the picture: the Bilbao site operates in a WRI 3.8/5 stress basin "
         "and the Company recycles only 12% of water withdrawal (sector best-practice: 45-60%). "
         "On a brighter note, EAF scrap utilisation at 71% exceeds the sector median."),
        tmpdir, chart_group="Climate")

    # ── Social ────────────────────────────────────────────────────────────
    _pillar_analysis(doc, "Social",
        ("The Social pillar score of 57/100 reflects an elevated occupational H&S profile and the "
         "complete absence of a Human Rights Policy or Human Rights Due Diligence framework. LTIFR "
         "of 2.8 per million hours is 33% above the sector median of 2.1, with one fatality in "
         "FY2022. ISO 45001 covers only 60% of sites. The absence of an SBTi target creates a "
         "distinct market-access risk: automotive OEMs (VW, Stellantis, BMW) imposing Scope 3 "
         "supplier decarbonisation requirements currently represent an estimated €456M (38% of "
         "revenue). Union density at 87% is a relative strength, as is training hours per employee."),
        tmpdir, chart_group="Workforce")

    # ── Governance ────────────────────────────────────────────────────────
    _pillar_analysis(doc, "Governance",
        ("The Governance pillar score of 49/100 is held back by two specific absences that are "
         "immediately remediable at low cost: no whistleblower mechanism (in breach of EU Directive "
         "2019/1937 since December 2021) and no ESG linkage in executive pay. Board independence "
         "at 78% (7 of 9) is above the peer median. CSRD readiness is the single largest "
         f"governance programme risk: as a large PIE, {D.COMPANY['short']} must publish a first ESRS-"
         "compliant report for FY2025 in Q1 2026, less than 18 months away. The double-materiality "
         "assessment is immature and the ESRS E1 transition plan is absent."),
        tmpdir, chart_group="Governance")

    # ── Climate Section ──────────────────────────────────────────────────
    _add_heading(doc, "Climate Risk Deep-Dive", level=1, page_break_before=True)
    _add_body(doc,
        "This section presents a TCFD-aligned climate risk assessment covering physical hazards, "
        "transition drivers, and forward-looking scenario analysis against three NGFS-calibrated "
        "pathways. All P&L figures are 2030 estimates relative to 2024 baseline.")
    _add_heading(doc, "Physical Risks", level=2)
    headers = ["Hazard", "Risk Score", "Affected Sites"]
    widths  = [5.0, 3.0, 7.0]
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    _add_table_header_row(tbl, headers, widths)
    for idx, pr in enumerate(cd.PHYSICAL_RISKS):
        sev = cd.severity(pr["score"])
        row = tbl.add_row()
        bg  = "F4F6F5" if idx % 2 == 0 else "FFFFFF"
        for i, (v, w) in enumerate(zip([pr["hazard"], f"{pr['score']}/100", pr["sites"]], widths)):
            cell = row.cells[i]
            cell.width = Cm(w)
            _cell_borders(cell, B["rule"])
            if i == 1:
                _cell_shade(cell, RAG_FILL.get(sev, bg))
                _cell_text(cell, v, bold=True, color=RAG_COLOR.get(sev, B["ink"]), center=True)
            else:
                _cell_shade(cell, bg)
                _cell_text(cell, v)
    doc.add_paragraph()

    _add_heading(doc, "Transition Risks", level=2)
    headers2 = ["Driver", "Risk Score"]
    widths2  = [10.0, 5.5]
    tbl2 = doc.add_table(rows=1, cols=2)
    tbl2.style = "Table Grid"
    _add_table_header_row(tbl2, headers2, widths2)
    for idx, tr in enumerate(cd.TRANSITION_RISKS):
        sev = cd.severity(tr["score"])
        row = tbl2.add_row()
        bg  = "F4F6F5" if idx % 2 == 0 else "FFFFFF"
        for i, (v, w) in enumerate(zip([tr["driver"], f"{tr['score']}/100"], widths2)):
            cell = row.cells[i]
            cell.width = Cm(w)
            _cell_borders(cell, B["rule"])
            if i == 1:
                _cell_shade(cell, RAG_FILL.get(sev, bg))
                _cell_text(cell, v, bold=True, color=RAG_COLOR.get(sev, B["ink"]), center=True)
            else:
                _cell_shade(cell, bg)
                _cell_text(cell, v)
    doc.add_paragraph()

    _add_heading(doc, "NGFS Scenario P&L Analysis (2030)", level=2)
    headers3 = ["NGFS Scenario", "Carbon Price 2030", "Estimated P&L Impact (€M)", "Key Note"]
    widths3  = [4.5, 3.0, 4.0, 4.0]
    tbl3 = doc.add_table(rows=1, cols=4)
    tbl3.style = "Table Grid"
    _add_table_header_row(tbl3, headers3, widths3)
    for idx, sc in enumerate(cd.CLIMATE_SCENARIOS):
        row = tbl3.add_row()
        bg  = "F4F6F5" if idx % 2 == 0 else "FFFFFF"
        for i, (v, w) in enumerate(zip(
            [sc["scenario"], sc["carbon_2030"], f"€{sc['pnl_2030_eur_m']}M", sc["note"]], widths3)):
            cell = row.cells[i]
            cell.width = Cm(w)
            _cell_borders(cell, B["rule"])
            _cell_shade(cell, bg)
            _cell_text(cell, v, bold=(i == 0), size=9)
    doc.add_paragraph()

    ch_climate = _chart_climate_scenarios(tmpdir)
    _insert_image(doc, ch_climate, width_cm=13,
                  caption="Figure 4: NGFS Scenario P&L Impact (2030) — all scenarios negative; Disorderly worst-case")

    # ── Risk Register ────────────────────────────────────────────────────
    _add_heading(doc, "Full Risk Register", level=1, page_break_before=True)
    _add_body(doc,
        f"The table below presents all {len(cd.RISKS)} identified ESG risks ranked by score. "
        "Severity classifications: Critical ≥75, High 55-74, Medium 35-54, Low <35.")
    _risk_register_table(doc)

    ch_sev = _chart_risk_severity(tmpdir)
    _insert_image(doc, ch_sev, width_cm=10,
                  caption="Figure 2: Risk Severity Distribution — majority High/Critical")

    # ── Financial Exposure ──────────────────────────────────────────────
    _add_heading(doc, "Quantified Financial Exposure", level=1, page_break_before=True)
    _add_body(doc,
        "The table and chart below summarise all quantified ESG financial exposures. "
        "Note: DRI/EAF transition capex (€500M phased) is a strategic investment, not a "
        "pure cost — it is the route to eliminating ETS and CBAM liabilities.")
    _financial_exposure_table(doc)
    ch_exp = _chart_financial_exposure(tmpdir)
    _insert_image(doc, ch_exp, width_cm=14,
                  caption="Figure 3: Quantified Financial Exposure by Driver (€M)")

    # ── Green Finance Opportunities ───────────────────────────────────────
    _add_heading(doc, "Green Finance Opportunities", level=1, page_break_before=True)
    _add_body(doc,
        "Addressing the three highest-priority ESG gaps would unlock substantial green-finance "
        "benefits. These are currently blocked primarily by the absence of an SBTi target and "
        "a published ESRS E1 climate transition plan.")
    _opportunities_table(doc)

    # ── Regulatory Timeline ──────────────────────────────────────────────
    _add_heading(doc, "Regulatory Timeline", level=1)
    _add_body(doc, "Key upcoming regulatory milestones with compliance status:")
    _reg_timeline_table(doc)

    # ── Engagement Asks ──────────────────────────────────────────────────
    _add_heading(doc, "Prioritised Engagement Asks", level=1, page_break_before=True)
    _add_body(doc,
        f"The following engagement asks are recommended to the Board of {D.COMPANY['short']} S.p.A. "
        "in order of priority. Critical and Immediate items should be commenced within 60 days.")
    _engagement_asks_table(doc)

    doc.save(out_path)


def _build_exec_docx(out_path: str, tmpdir: str) -> None:
    doc = _new_doc()
    _add_footer(doc)
    _cover_page(doc, cd.REPORT_TYPES["exec"],
                subtitle=f"2-Page Summary  |  {cd.COMPANY['report_period']}")
    _scorecard_section(doc)
    ch_pillars = _chart_pillar_scores(tmpdir)
    _insert_image(doc, ch_pillars, width_cm=12,
                  caption="Pillar Risk Scores")
    _exec_summary_text(doc)
    _add_heading(doc, "Top 5 Risks", level=2)
    _risk_register_table(doc, top_n=5)
    _add_heading(doc, "Quantified Financial Exposure", level=2)
    _financial_exposure_table(doc)
    ch_exp = _chart_financial_exposure(tmpdir)
    _insert_image(doc, ch_exp, width_cm=12, caption="Financial Exposure (€M)")
    _add_heading(doc, "Priority Engagement Asks", level=2)
    _engagement_asks_table(doc)
    doc.save(out_path)


def _build_climate_docx(out_path: str, tmpdir: str) -> None:
    doc = _new_doc()
    _add_footer(doc)
    _cover_page(doc, cd.REPORT_TYPES["climate"],
                subtitle=f"TCFD-Aligned Climate Risk Assessment  |  {cd.COMPANY['report_period']}")

    # Governance
    _add_heading(doc, "1. Governance", level=1)
    _add_body(doc,
        f"{D.COMPANY['short']}'s Board retains ultimate oversight of climate risk through its Risk Committee, "
        "which meets quarterly. However, no Board-level climate expertise is formally disclosed, "
        "the Chief Sustainability Officer has no direct Board reporting line, and climate performance "
        "is not linked to executive remuneration. These structural gaps limit the Board's ability to "
        "respond decisively to the material climate exposures outlined below.")

    # Strategy
    _add_heading(doc, "2. Strategy", level=1)
    _add_body(doc,
        f"{D.COMPANY['short']}'s 100% BF-BOF primary steel route generates 842 ktCO2e Scope 1 per annum "
        "— a carbon intensity of approximately 0.73 tCO2e per €k revenue. The Company has not "
        "published a climate transition plan, committed to an SBTi pathway, or set renewable energy "
        "targets. Against the backdrop of three materially different climate scenarios (see Section 4), "
        "the current strategy is insufficiently resilient under any pathway.")
    _add_heading(doc, "Physical Risks", level=2)
    headers = ["Hazard", "Risk Score", "Sites Affected"]
    widths  = [5.0, 3.0, 7.5]
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    _add_table_header_row(tbl, headers, widths)
    for idx, pr in enumerate(cd.PHYSICAL_RISKS):
        sev = cd.severity(pr["score"])
        row = tbl.add_row()
        bg  = "F4F6F5" if idx % 2 == 0 else "FFFFFF"
        for i, (v, w) in enumerate(zip([pr["hazard"], f"{pr['score']}/100", pr["sites"]], widths)):
            cell = row.cells[i]
            cell.width = Cm(w)
            _cell_borders(cell, B["rule"])
            if i == 1:
                _cell_shade(cell, RAG_FILL.get(sev, bg))
                _cell_text(cell, v, bold=True, color=RAG_COLOR.get(sev, B["ink"]), center=True)
            else:
                _cell_shade(cell, bg)
                _cell_text(cell, v)
    doc.add_paragraph()
    _add_heading(doc, "Transition Risks", level=2)
    headers2 = ["Driver", "Risk Score", "Severity"]
    widths2  = [7.0, 3.0, 3.0]
    tbl2 = doc.add_table(rows=1, cols=3)
    tbl2.style = "Table Grid"
    _add_table_header_row(tbl2, headers2, widths2)
    for idx, tr in enumerate(cd.TRANSITION_RISKS):
        sev = cd.severity(tr["score"])
        row = tbl2.add_row()
        bg  = "F4F6F5" if idx % 2 == 0 else "FFFFFF"
        for i, (v, w) in enumerate(zip([tr["driver"], f"{tr['score']}/100", sev], widths2)):
            cell = row.cells[i]
            cell.width = Cm(w)
            _cell_borders(cell, B["rule"])
            if i in (1, 2):
                _cell_shade(cell, RAG_FILL.get(sev, bg))
                _cell_text(cell, v, bold=True, color=RAG_COLOR.get(sev, B["ink"]), center=True)
            else:
                _cell_shade(cell, bg)
                _cell_text(cell, v)
    doc.add_paragraph()

    # Risk Management
    _add_heading(doc, "3. Risk Management", level=1)
    _add_body(doc,
        "No formal climate risk identification and assessment process is documented in public "
        "disclosures. Material climate risks — EU ETS (ENV-01, score 89), CBAM (ENV-02, score 84), "
        "and technology/stranding risk (ENV-03, score 76) — are identified in our analysis but "
        "not managed against documented thresholds. No ETS hedging policy, no internal carbon "
        "price, and no Scope 3 mapping beyond Category 1 (purchased goods) are in place.")
    _add_heading(doc, "ETS and CBAM Deep-Dive", level=2)
    _add_body(doc,
        "EU ETS Phase 4 free allocation tapers from approximately 40% in 2024 to 10% by 2030. "
        "At 842 ktCO2e Scope 1 and a 2030 ETS price of €90/t (Orderly scenario), the net annual "
        "cash cost reaches €68M. Under CBAM (full phase-in 2026), embedded carbon in roughly "
        "€280M of non-EU exports (at ~0.85 tCO2e/t) generates €18-28M additional certificate "
        "cost annually. These liabilities are only structurally reduced by decarbonising the "
        "production route — process efficiency alone is insufficient at this carbon intensity.")

    # Metrics & Targets
    _add_heading(doc, "4. Metrics & Targets", level=1)
    _add_heading(doc, "NGFS Scenario P&L Analysis (2030)", level=2)
    _add_body(doc, "All figures are estimated incremental P&L vs 2024 baseline, attributable to climate factors.")
    headers3 = ["NGFS Scenario", "Carbon Price 2030", "P&L Impact (€M)", "Key Note"]
    widths3  = [4.5, 3.0, 3.5, 4.5]
    tbl3 = doc.add_table(rows=1, cols=4)
    tbl3.style = "Table Grid"
    _add_table_header_row(tbl3, headers3, widths3)
    for idx, sc in enumerate(cd.CLIMATE_SCENARIOS):
        row = tbl3.add_row()
        bg  = "F4F6F5" if idx % 2 == 0 else "FFFFFF"
        for i, (v, w) in enumerate(zip(
            [sc["scenario"], sc["carbon_2030"], f"€{sc['pnl_2030_eur_m']}M", sc["note"]], widths3)):
            cell = row.cells[i]
            cell.width = Cm(w)
            _cell_borders(cell, B["rule"])
            _cell_shade(cell, bg)
            _cell_text(cell, v, bold=(i == 0))
    doc.add_paragraph()
    ch_climate = _chart_climate_scenarios(tmpdir)
    _insert_image(doc, ch_climate, width_cm=13,
                  caption="NGFS Scenario P&L Impact (2030) — all scenarios result in material negative P&L")

    ch_kpi = _chart_kpi_heatmap(tmpdir, group="Climate")
    _insert_image(doc, ch_kpi, width_cm=13,
                  caption="Climate KPI Peer Percentiles — Scope 1 at P89 is bottom-decile vs peers")

    # Recommendations
    _add_heading(doc, "5. Recommendations", level=1)
    climate_asks = [a for a in cd.ENGAGEMENT_ASKS
                    if any(kw in a["ask"] for kw in ["SBTi", "climate", "ESRS E1", "whistleblow"])]
    for ask in climate_asks[:4]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"[{ask['priority']} — {ask['by']}] ")
        r.font.bold = True; r.font.color.rgb = _rgb("accent"); r.font.size = Pt(9)
        r2 = p.add_run(ask["ask"])
        r2.font.size = Pt(9); r2.font.color.rgb = _rgb("slate")

    doc.save(out_path)


def _build_engagement_docx(out_path: str, tmpdir: str) -> None:
    doc = _new_doc()
    _add_footer(doc)

    # Letterhead
    p = doc.add_paragraph()
    r = p.add_run(B["platform"].upper() + "  |  ESG ENGAGEMENT LETTER")
    r.font.size = Pt(10); r.font.bold = True
    r.font.color.rgb = _rgb("accent"); r.font.name = B["font_body"]
    _set_para_fmt(p, space_after=2)
    phr = doc.add_paragraph()
    pPr = phr._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "8")
    bot.set(qn("w:space"), "1");   bot.set(qn("w:color"), B["accent"])
    pBdr.append(bot); pPr.append(pBdr)
    _set_para_fmt(phr, space_after=16)

    # Date and addressee
    _add_body(doc, cd.COMPANY["as_of"], size=10, space_after=10)
    _add_body(doc, "To: The Board of Directors", bold=True, size=10, space_after=2)
    _add_body(doc, f"{cd.COMPANY['name']}", size=10, space_after=2)
    _add_body(doc, f"{cd.COMPANY['domicile']}", size=10, space_after=16)

    # Subject
    p_subj = doc.add_paragraph()
    r1 = p_subj.add_run("Subject: ")
    r1.font.bold = True; r1.font.size = Pt(11); r1.font.name = B["font_body"]; r1.font.color.rgb = _rgb("ink")
    r2 = p_subj.add_run(f"ESG Engagement — {cd.COMPANY['name']} ({cd.COMPANY['report_period']})")
    r2.font.size = Pt(11); r2.font.name = B["font_body"]; r2.font.color.rgb = _rgb("ink")
    _set_para_fmt(p_subj, space_after=14)

    # Body
    paras = [
        ("Dear Members of the Board,",),
        ("We write on behalf of our investment fund as a shareholder in "
         f"{cd.COMPANY['name']} (\"the Company\") following completion of our annual ESG due "
         f"diligence review for {cd.COMPANY['report_period']}. We engage constructively with "
         "portfolio companies on material ESG risks and view this letter as the beginning of "
         "a structured dialogue aimed at protecting and enhancing long-term enterprise value.",),
        ("CONTEXT AND MATERIALITY",),
        (f"Our proprietary ESG risk assessment places the Company at {cd.SCORES['overall']}/100 "
         f"(High Risk, P{cd.SCORES['overall_percentile']} vs NACE C24 peers). This rating is "
         "driven by three material risk clusters: (i) climate and carbon transition costs "
         f"(estimated €{cd.FINANCIAL_EXPOSURE[0]['amount_eur_m'] + cd.FINANCIAL_EXPOSURE[1]['amount_eur_m']}M combined "
         "annual exposure by 2027 from EU ETS and CBAM); (ii) human rights and social compliance "
         "gaps that expose the Company to CSDDD fines of up to €60M; and (iii) governance "
         "deficiencies — specifically the absence of a whistleblower mechanism and CSRD readiness "
         "gaps — that represent both regulatory non-compliance and reputational risk.",),
        ("INVESTMENT IMPLICATIONS",),
        ("The Company's current carbon profile (842 ktCO2e Scope 1, BF-BOF route, no SBTi "
         "target) is increasingly incompatible with the sustainability commitments of key "
         f"automotive customers. We estimate approximately €{cd.RISKS[9]['financial']} of revenue "
         "is at risk as OEM Scope 3 supplier requirements tighten. Green-bond and sustainability-"
         "linked loan access — representing €400-550M of potential lower-cost financing — is "
         "currently blocked by the absence of an SBTi-validated target.",),
        ("ENGAGEMENT ASKS",),
        ("We request that the Board address the following items, detailed in the table below, "
         "within the stated timeframes. We request written responses confirming the Company's "
         "position on each item within 30 days of receipt of this letter.",),
    ]
    italic_heads = {"CONTEXT AND MATERIALITY", "INVESTMENT IMPLICATIONS", "ENGAGEMENT ASKS"}
    for content in paras:
        text = content[0]
        if text in italic_heads:
            _add_heading(doc, text, level=2, space_before=10, space_after=4)
        else:
            _add_body(doc, text, size=10, space_after=8)

    # Engagement asks table
    _engagement_asks_table(doc)

    ch_pillars = _chart_pillar_scores(tmpdir)
    _insert_image(doc, ch_pillars, width_cm=12,
                  caption="ESG Pillar Risk Scores — Environmental is primary concern at 64/100")

    # Closing
    _add_body(doc,
        "We welcome the opportunity to discuss these matters directly with the Board and "
        "management team. We propose an initial engagement meeting within 45 days. Please "
        "direct responses and scheduling to ESGIntel's engagement team.", size=10, space_after=10)
    _add_body(doc, "Yours faithfully,", size=10, space_after=16)
    _add_body(doc, "ESGIntel ESG Research & Engagement", bold=True, size=10, space_after=2)
    _add_body(doc, B["confidential"], italic=True, size=8, color_key="muted", space_after=0)

    doc.save(out_path)


def _build_riskregister_docx(out_path: str, tmpdir: str) -> None:
    doc = _new_doc()
    _add_footer(doc)
    _cover_page(doc, cd.REPORT_TYPES["riskregister"],
                subtitle=f"Full Risk Register  |  {cd.COMPANY['report_period']}")

    # Summary distribution
    _add_heading(doc, "Severity Distribution Summary", level=1)
    from collections import Counter
    counts = Counter(cd.severity(r["score"]) for r in cd.RISKS)
    total  = len(cd.RISKS)
    _add_body(doc, f"Total risks identified: {total}", bold=True, size=10)
    for sev in ["Critical", "High", "Medium", "Low"]:
        n = counts.get(sev, 0)
        p = doc.add_paragraph()
        r = p.add_run(f"  {sev}: {n}  ({n/total*100:.0f}%)")
        r.font.size = Pt(10)
        r.font.bold = sev in ("Critical", "High")
        r.font.color.rgb = RGBColor(*cd.hex_rgb(RAG_COLOR.get(sev, B["ink"])))
        r.font.name = B["font_body"]
        _set_para_fmt(p, space_after=2)

    ch_sev = _chart_risk_severity(tmpdir)
    _insert_image(doc, ch_sev, width_cm=10,
                  caption="Risk Severity Distribution")

    # Full table
    _add_heading(doc, "Full Risk Register", level=1)
    _add_body(doc,
        "Sorted by risk score descending. Severity: Critical≥75, High 55-74, Medium 35-54, Low<35.")
    sorted_risks = sorted(cd.RISKS, key=lambda r: r["score"], reverse=True)
    _risk_register_table(doc, risks=sorted_risks)

    ch_exp = _chart_financial_exposure(tmpdir)
    _insert_image(doc, ch_exp, width_cm=14,
                  caption="Quantified Financial Exposure by Risk Driver (€M)")

    # Detail cards
    _add_heading(doc, "Risk Detail Narratives", level=1, page_break_before=True)
    for risk in sorted_risks:
        sev = cd.severity(risk["score"])
        _add_heading(doc, f"{risk['id']} — {risk['title']}", level=2, space_before=10)
        p = doc.add_paragraph()
        r1 = p.add_run(f"Severity: {sev}  |  Score: {risk['score']}/100  |  Horizon: {risk['horizon']}")
        r1.font.size = Pt(9); r1.font.bold = True
        r1.font.color.rgb = RGBColor(*cd.hex_rgb(RAG_COLOR.get(sev, B["ink"])))
        _set_para_fmt(p, space_after=3)
        _add_body(doc, risk["summary"], size=9, space_after=3)
        p2 = doc.add_paragraph()
        r2a = p2.add_run("Financial Exposure: ")
        r2a.font.bold = True; r2a.font.size = Pt(9); r2a.font.color.rgb = _rgb("ink")
        r2b = p2.add_run(risk["financial"])
        r2b.font.size = Pt(9); r2b.font.color.rgb = _rgb("slate")
        _set_para_fmt(p2, space_after=3)
        p3 = doc.add_paragraph()
        r3a = p3.add_run("Recommended Action: ")
        r3a.font.bold = True; r3a.font.size = Pt(9); r3a.font.color.rgb = _rgb("accent")
        r3b = p3.add_run(risk["action"])
        r3b.font.size = Pt(9); r3b.font.color.rgb = _rgb("slate")
        _set_para_fmt(p3, space_after=8)

    doc.save(out_path)


# ════════════════════════════════════════════════════════════════════════════
# PDF BUILDER (reportlab)
# ════════════════════════════════════════════════════════════════════════════

def _rl_color(hex_key: str):
    r, g, b = cd.hex_rgb(B[hex_key])
    return colors.Color(r/255, g/255, b/255)

def _rl_hex(hex_str: str):
    r, g, b = cd.hex_rgb(hex_str)
    return colors.Color(r/255, g/255, b/255)


def _build_pdf(report_type: str, out_path: str, tmpdir: str) -> None:
    """Generic PDF renderer using ReportLab."""
    title = cd.REPORT_TYPES.get(report_type, "ESG Report")
    doc = SimpleDocTemplate(
        out_path,
        pagesize=A4,
        leftMargin=2.2*cm, rightMargin=2.2*cm,
        topMargin=2.0*cm,  bottomMargin=2.5*cm,
        title=f"{cd.COMPANY['name']} — {title}",
        author=B["platform"],
    )

    # Style sheet
    SS = getSampleStyleSheet()
    ink   = _rl_color("ink")
    acc   = _rl_color("accent")
    slate = _rl_color("slate")
    muted = _rl_color("muted")
    red_c = _rl_color("red")
    page_w = A4[0] - 4.4*cm  # content width

    def _s(name, **kw):
        base = SS["Normal"]
        return ParagraphStyle(name, parent=base, **kw)

    ST = {
        "h1":    _s("h1",    fontSize=16, textColor=ink, fontName="Helvetica-Bold",
                             spaceAfter=6, spaceBefore=14, leading=20),
        "h2":    _s("h2",    fontSize=12, textColor=ink, fontName="Helvetica-Bold",
                             spaceAfter=4, spaceBefore=10, leading=15),
        "h3":    _s("h3",    fontSize=10, textColor=acc, fontName="Helvetica-Bold",
                             spaceAfter=3, spaceBefore=6, leading=13),
        "body":  _s("body",  fontSize=9,  textColor=slate, fontName="Helvetica",
                             spaceAfter=5, leading=13),
        "small": _s("small", fontSize=7.5, textColor=muted, fontName="Helvetica-Oblique",
                             spaceAfter=3, leading=10),
        "cover_title": _s("ct", fontSize=24, textColor=ink, fontName="Helvetica-Bold",
                           alignment=TA_CENTER, spaceAfter=8, leading=30),
        "cover_sub":   _s("cs", fontSize=13, textColor=acc, fontName="Helvetica-Bold",
                           alignment=TA_CENTER, spaceAfter=6),
        "cover_meta":  _s("cm", fontSize=9,  textColor=slate, fontName="Helvetica",
                           alignment=TA_CENTER, spaceAfter=3),
        "conf":        _s("cf", fontSize=8,  textColor=muted, fontName="Helvetica-Oblique",
                           alignment=TA_CENTER, spaceAfter=4),
        "bullet": _s("blt",  fontSize=9, textColor=slate, fontName="Helvetica",
                             spaceAfter=4, leftIndent=16, bulletIndent=8,
                             bulletFontName="Helvetica", leading=13),
    }

    rag_rl = {
        "Critical": (colors.HexColor("#FAD7D4"), colors.HexColor("#C0392B")),
        "High":     (colors.HexColor("#FDEBD0"), colors.HexColor("#E67E22")),
        "Medium":   (colors.HexColor("#FEF9E7"), colors.HexColor("#C77D11")),
        "Low":      (colors.HexColor("#D5F5E3"), colors.HexColor("#1E8E5A")),
    }

    def tbl_style(header_rows=1) -> list:
        return [
            ("BACKGROUND",   (0,0), (-1, header_rows-1), _rl_color("ink")),
            ("TEXTCOLOR",    (0,0), (-1, header_rows-1), colors.white),
            ("FONTNAME",     (0,0), (-1, header_rows-1), "Helvetica-Bold"),
            ("FONTSIZE",     (0,0), (-1,-1), 8),
            ("FONTNAME",     (0, header_rows), (-1,-1), "Helvetica"),
            ("TEXTCOLOR",    (0, header_rows), (-1,-1), _rl_color("slate")),
            ("ROWBACKGROUNDS",(0, header_rows),(-1,-1),
             [_rl_color("row_alt"), colors.white]),
            ("GRID",         (0,0), (-1,-1), 0.3, _rl_color("rule")),
            ("ALIGN",        (0,0), (-1,-1), "LEFT"),
            ("VALIGN",       (0,0), (-1,-1), "MIDDLE"),
            ("TOPPADDING",   (0,0), (-1,-1), 4),
            ("BOTTOMPADDING",(0,0), (-1,-1), 4),
            ("LEFTPADDING",  (0,0), (-1,-1), 5),
            ("RIGHTPADDING", (0,0), (-1,-1), 5),
        ]

    story = []

    # ── Cover ──────────────────────────────────────────────────────────────
    story.append(Spacer(1, 3*cm))
    story.append(Paragraph(B["platform"].upper(), ST["h3"]))
    story.append(Paragraph(B["tagline"], ST["small"]))
    story.append(HRFlowable(width=page_w, thickness=2, color=acc, spaceAfter=12))
    story.append(Paragraph(cd.COMPANY["name"], ST["cover_title"]))
    story.append(Paragraph(cd.COMPANY["sector"], ST["cover_meta"]))
    story.append(Spacer(1, 0.4*cm))
    story.append(Paragraph(title, ST["cover_sub"]))
    story.append(Spacer(1, 0.5*cm))
    for lbl, val in [("Period", cd.COMPANY["report_period"]),
                     ("As of",  cd.COMPANY["as_of"]),
                     ("Domicile", cd.COMPANY["domicile"])]:
        story.append(Paragraph(f"<b>{lbl}:</b>  {val}", ST["cover_meta"]))
    story.append(Spacer(1, 1*cm))
    story.append(Paragraph(B["confidential"].upper(), ST["conf"]))
    story.append(PageBreak())

    # ── Scorecard ─────────────────────────────────────────────────────────
    story.append(Paragraph("ESG Risk Scorecard", ST["h1"]))
    story.append(Paragraph(
        f"<b>Overall Score: {cd.SCORES['overall']}/100 — {cd.SCORES['overall_label']} "
        f"(P{cd.SCORES['overall_percentile']} vs peers)</b>", ST["h2"]))
    sc_data = [["Pillar", "Score", "Level", "Peer Pctl", "Weight"]]
    for pillar, data in cd.SCORES["pillars"].items():
        sc_data.append([pillar, f"{data['score']}/100", data["label"],
                         f"P{data['percentile']}", f"{int(data['weight']*100)}%"])
    sc_tbl = Table(sc_data, colWidths=[4.5*cm, 2.5*cm, 2.5*cm, 3*cm, 2*cm])
    sc_style = tbl_style()
    for i, (_, data) in enumerate(cd.SCORES["pillars"].items(), 1):
        bg, tc = rag_rl.get(data["label"], (colors.white, _rl_color("ink")))
        sc_style.append(("BACKGROUND", (2,i), (2,i), bg))
        sc_style.append(("TEXTCOLOR",  (2,i), (2,i), tc))
        sc_style.append(("FONTNAME",   (2,i), (2,i), "Helvetica-Bold"))
    sc_tbl.setStyle(TableStyle(sc_style))
    story.append(sc_tbl)
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph(f"Methodology: {cd.SCORES['formula']}", ST["small"]))
    story.append(Spacer(1, 0.3*cm))

    # Pillar chart
    ch = _chart_pillar_scores(tmpdir)
    if ch and os.path.exists(ch):
        story.append(RLImage(ch, width=12*cm, height=5.5*cm))
        story.append(Paragraph("Figure 1: Pillar Risk Scores", ST["small"]))

    # ── Report-specific content ────────────────────────────────────────────
    if report_type in ("full", "exec"):
        story.append(PageBreak())
        story.append(Paragraph("Executive Summary", ST["h1"]))
        story.append(Paragraph(
            f"{cd.COMPANY['name']} presents a HIGH ESG risk profile ({cd.SCORES['overall']}/100, "
            f"P{cd.SCORES['overall_percentile']}). Three risk clusters dominate: climate/carbon "
            f"(€{cd.FINANCIAL_EXPOSURE[0]['amount_eur_m']+cd.FINANCIAL_EXPOSURE[1]['amount_eur_m']}M combined "
            "ETS+CBAM exposure by 2027), human-rights/compliance (CSDDD €60M fine risk), and "
            "market-access risk (€456M automotive revenue at risk from absent SBTi).", ST["body"]))

    if report_type == "full":
        story.append(Spacer(1, 0.3*cm))
        # KPI tables per pillar
        for pillar, groups in [
            ("Environmental", ["Climate (E1)", "Water (E3)", "Circular (E5)"]),
            ("Social",        ["Workforce (S1)"]),
            ("Governance",    ["Governance (G1)"]),
        ]:
            story.append(PageBreak())
            story.append(Paragraph(f"{pillar} Pillar KPIs", ST["h1"]))
            kpis = [k for k in cd.KPIS if k["group"] in groups and k["pctl"]]
            kpi_data = [["Metric", "Group", "Value", "Peer", "Pctl"]]
            for k in kpis:
                kpi_data.append([k["metric"], k["group"], k["value"],
                                  k["peer"] or "—", f"P{k['pctl']}"])
            kpi_tbl = Table(kpi_data,
                            colWidths=[5*cm, 3*cm, 3*cm, 2.5*cm, 1.5*cm])
            kpi_tbl.setStyle(TableStyle(tbl_style()))
            story.append(kpi_tbl)
            ch_kpi = _chart_kpi_heatmap(tmpdir, group=groups[0].split("(")[0].strip())
            if ch_kpi and os.path.exists(ch_kpi):
                story.append(Spacer(1, 0.2*cm))
                story.append(RLImage(ch_kpi, width=13*cm, height=5*cm))
                story.append(Paragraph(f"Figure: {pillar} KPI Peer Percentiles", ST["small"]))

    # Risk register
    story.append(PageBreak())
    story.append(Paragraph("Risk Register", ST["h1"]))
    top_n = 5 if report_type == "exec" else None
    risks = sorted(cd.RISKS, key=lambda r: r["score"], reverse=True)
    if top_n:
        risks = risks[:top_n]
    rr_data = [["ID", "Risk", "Sev", "Score", "Horizon", "Financial"]]
    rr_style = tbl_style()
    for i, risk in enumerate(risks, 1):
        sev = cd.severity(risk["score"])
        bg, tc = rag_rl.get(sev, (colors.white, _rl_color("ink")))
        rr_data.append([risk["id"], Paragraph(risk["title"], ST["body"]),
                         sev, str(risk["score"]), risk["horizon"], risk["financial"]])
        rr_style.append(("BACKGROUND", (2,i), (2,i), bg))
        rr_style.append(("TEXTCOLOR",  (2,i), (2,i), tc))
        rr_style.append(("FONTNAME",   (2,i), (2,i), "Helvetica-Bold"))
    rr_tbl = Table(rr_data, colWidths=[1.5*cm, 4*cm, 1.8*cm, 1.3*cm, 2.5*cm, 3.9*cm])
    rr_tbl.setStyle(TableStyle(rr_style))
    story.append(rr_tbl)

    ch_sev = _chart_risk_severity(tmpdir)
    if ch_sev and os.path.exists(ch_sev):
        story.append(Spacer(1, 0.3*cm))
        story.append(RLImage(ch_sev, width=9*cm, height=7*cm))
        story.append(Paragraph("Risk Severity Distribution", ST["small"]))

    # Financial exposure
    story.append(PageBreak())
    story.append(Paragraph("Quantified Financial Exposure", ST["h1"]))
    fe_data = [["Driver", "Exposure", "Type", "Rating"]]
    fe_style = tbl_style()
    for i, fe in enumerate(cd.FINANCIAL_EXPOSURE, 1):
        sev = fe["rating"]
        bg, tc = rag_rl.get(sev, (colors.white, _rl_color("ink")))
        fe_data.append([Paragraph(fe["driver"], ST["body"]),
                         f"€{fe['amount_eur_m']}M", fe["type"], fe["rating"]])
        fe_style.append(("BACKGROUND", (3,i), (3,i), bg))
        fe_style.append(("TEXTCOLOR",  (3,i), (3,i), tc))
        fe_style.append(("FONTNAME",   (3,i), (3,i), "Helvetica-Bold"))
    fe_tbl = Table(fe_data, colWidths=[6.5*cm, 2.5*cm, 3.5*cm, 2.5*cm])
    fe_tbl.setStyle(TableStyle(fe_style))
    story.append(fe_tbl)
    ch_exp = _chart_financial_exposure(tmpdir)
    if ch_exp and os.path.exists(ch_exp):
        story.append(Spacer(1, 0.3*cm))
        story.append(RLImage(ch_exp, width=13*cm, height=5*cm))
        story.append(Paragraph("Financial Exposure by Driver (€M)", ST["small"]))

    if report_type in ("full", "climate"):
        story.append(PageBreak())
        story.append(Paragraph("NGFS Climate Scenario Analysis", ST["h1"]))
        sc_data2 = [["Scenario", "Carbon Price", "P&L Impact (€M)", "Note"]]
        for sc in cd.CLIMATE_SCENARIOS:
            sc_data2.append([sc["scenario"], sc["carbon_2030"],
                              f"€{sc['pnl_2030_eur_m']}M", sc["note"]])
        sc_tbl2 = Table(sc_data2, colWidths=[5*cm, 3*cm, 3.5*cm, 3.5*cm])
        sc_tbl2.setStyle(TableStyle(tbl_style()))
        story.append(sc_tbl2)
        ch_cl = _chart_climate_scenarios(tmpdir)
        if ch_cl and os.path.exists(ch_cl):
            story.append(Spacer(1, 0.3*cm))
            story.append(RLImage(ch_cl, width=12*cm, height=5*cm))
            story.append(Paragraph("NGFS Scenario P&L Impact (2030)", ST["small"]))

    # Opportunities & Engagement
    story.append(PageBreak())
    story.append(Paragraph("Green Finance Opportunities", ST["h1"]))
    opp_data = [["Lever", "Size", "Benefit", "Pre-condition"]]
    for o in cd.OPPORTUNITIES:
        opp_data.append([o["lever"], o["size"], o["benefit"], o["precondition"]])
    opp_tbl = Table(opp_data, colWidths=[4*cm, 2.5*cm, 4*cm, 4.5*cm])
    opp_tbl.setStyle(TableStyle(tbl_style()))
    story.append(opp_tbl)

    story.append(Spacer(1, 0.5*cm))
    story.append(Paragraph("Engagement Asks", ST["h1"]))
    ea_data = [["Ask", "Owner", "By", "Priority"]]
    ea_style = tbl_style()
    for i, ask in enumerate(cd.ENGAGEMENT_ASKS, 1):
        prio = ask["priority"]
        bg, tc = rag_rl.get(prio, (colors.white, _rl_color("ink")))
        ea_data.append([Paragraph(ask["ask"], ST["body"]),
                         ask["owner"], ask["by"], prio])
        ea_style.append(("BACKGROUND", (3,i), (3,i), bg))
        ea_style.append(("TEXTCOLOR",  (3,i), (3,i), tc))
        ea_style.append(("FONTNAME",   (3,i), (3,i), "Helvetica-Bold"))
    ea_tbl = Table(ea_data, colWidths=[7*cm, 3*cm, 2.5*cm, 2.5*cm])
    ea_tbl.setStyle(TableStyle(ea_style))
    story.append(ea_tbl)

    # Regulatory timeline
    if report_type in ("full", "climate"):
        story.append(PageBreak())
        story.append(Paragraph("Regulatory Timeline", ST["h1"]))
        rt_data = [["Year", "Item", "Status"]]
        for item in cd.REG_TIMELINE:
            rt_data.append([item["year"], item["item"], item["status"]])
        rt_tbl = Table(rt_data, colWidths=[1.8*cm, 10*cm, 3.2*cm])
        rt_tbl.setStyle(TableStyle(tbl_style()))
        story.append(rt_tbl)

    # Footer page number via canvas
    def _on_page(canvas, doc_obj):
        canvas.saveState()
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(muted)
        w, h = A4
        canvas.drawCentredString(
            w/2, 1.2*cm,
            f"{B['platform']} — CONFIDENTIAL  |  {cd.COMPANY['name']}  |  {cd.COMPANY['report_period']}  |  Page {doc_obj.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=_on_page, onLaterPages=_on_page)


# ════════════════════════════════════════════════════════════════════════════
# PUBLIC ENTRY POINT
# ════════════════════════════════════════════════════════════════════════════

def build_report(report_type: str, fmt: str, out_dir: str) -> str:
    """
    Build an ESG report document.

    Parameters
    ----------
    report_type : str
        One of 'full', 'exec', 'climate', 'engagement', 'riskregister'.
    fmt : str
        'docx' or 'pdf'.
    out_dir : str
        Directory where the output file will be written. Created if absent.

    Returns
    -------
    str
        Absolute path to the generated file.
    """
    if report_type not in cd.REPORT_TYPES:
        raise ValueError(f"Unknown report_type '{report_type}'. "
                         f"Choose from: {list(cd.REPORT_TYPES)}")
    if fmt not in ("docx", "pdf"):
        raise ValueError(f"Unknown fmt '{fmt}'. Choose 'docx' or 'pdf'.")

    out_dir = os.path.abspath(out_dir)
    os.makedirs(out_dir, exist_ok=True)

    title   = cd.REPORT_TYPES[report_type]
    safe_t  = cd.slug(title).replace("_", " ").title().replace(" ", "_")
    fname   = f"{D.COMPANY['short'].replace(' ', '_')}_{safe_t}.{fmt}"
    out_path = os.path.join(out_dir, fname)

    tmpdir = tempfile.mkdtemp(prefix="esgIntel_charts_")
    try:
        if fmt == "docx":
            builders = {
                "full":         _build_full_docx,
                "exec":         _build_exec_docx,
                "climate":      _build_climate_docx,
                "engagement":   _build_engagement_docx,
                "riskregister": _build_riskregister_docx,
            }
            builders[report_type](out_path, tmpdir)
        else:
            _build_pdf(report_type, out_path, tmpdir)
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)

    return out_path
